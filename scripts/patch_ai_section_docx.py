"""
Insert or refresh section 5.2.3.1 (AI architecture) in an existing SIGTS final-project .docx.
Use when you maintain a hand-edited report (e.g. FINAL PROJECT REVISED - 2.docx).

  python scripts/patch_ai_section_docx.py
  python scripts/patch_ai_section_docx.py --path "c:\\path\\to\\your.docx"
"""
from __future__ import annotations

import argparse
import sys
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DEFAULT_PATH = r"c:\Projects\SIGTS\FINAL PROJECT REVISED - 2.docx"

MARKER_5231 = "5.2.3.1 Artificial Intelligence Architecture And Algorithms"
MARKER_524 = "5.2.4 Guide And IT Operations"
MARKER_523 = "5.2.3 Tour Help And Catalogues"

TOUR_HELP_INTRO = (
    "The system allowed tourists to browse species and cultural catalogues and to receive tour-help "
    "responses through a conversational module (“Tour Help”) integrated in the progressive web "
    "application. Responses were not produced by a proprietary machine-learning model trained "
    "within the project; instead, the implementation combined (i) an optional hosted large "
    "language model (LLM) accessed through an OpenAI-compatible application programming interface, "
    "and (ii) a rule-based knowledge interpreter used when no API key was configured, when the "
    "network or provider failed, or when questions were filtered as clearly off-topic."
)

AI_PARAS = [
    (
        "Table 5.2 summarises the technical stack. The primary generative path, when enabled, used "
        "the Chat Completions protocol (POST /v1/chat/completions) with default model GPT-4o, "
        "configurable through environment variables (SIGTS_CHAT_OPENAI_API_KEY, SIGTS_CHAT_MODEL, "
        "SIGTS_CHAT_TEMPERATURE approximately 0.52, SIGTS_CHAT_MAX_TOKENS up to 1024). GPT-4o is "
        "a commercial transformer-based decoder-only language model; SIGTS invoked it as a remote "
        "service and did not implement neural training or inference locally."
    ),
    (
        "Before each LLM call, the backend performed retrieval-augmented generation in a lightweight "
        "form (RAG-lite). The visitor question was tokenised (lowercase words of at least three "
        "characters), converted to SQL ILIKE patterns, and matched against PostgreSQL tables "
        "including parks, FAQs, safety tips, destination information, animals, wildlife tour themes, "
        "cultural narratives, locations, tour routes, and recent sightings. A core Bwindi briefing "
        "block was always appended so the model retained baseline park facts even when lexical "
        "matches were weak. The client also supplied a sanitised catalogue snapshot (themes, animals, "
        "map points, stories, FAQs, safety species) assembled in the browser. Retrieved text was "
        "injected into a system prompt that instructed the model to stay on Bwindi visitor topics, "
        "paraphrase grounded facts, and decline to invent permit prices or citations. This approach "
        "did not use vector embeddings or a semantic search index; matching was keyword-based only."
    ),
    (
        "The fallback algorithm (modes labelled rule_kb_v1 and rule_kb_v1_fallback) was a "
        "deterministic expert-system style pipeline: regular-expression intent detection for themes "
        "such as gorillas, safety, weather, culture, and map routing; scope checks for Bwindi and "
        "nature-tourism vocabulary; substring matching against the on-device species catalogue; and "
        "template paragraphs for greetings and general redirection. The same rule order was "
        "mirrored in JavaScript for offline use when POST /api/ai/chat was unavailable. Eligible "
        "questions were logged in the ai_query_logs table with response time and language metadata. "
        "Configuration could be inspected at GET /api/ai/status (llm_configured, model, provider, "
        "grounding tables)."
    ),
    (
        "The request flow was as follows: an authenticated tourist submitted a question; "
        "optional GPS coordinates resolve to the nearest map point of interest, retrieval builds the "
        "knowledge pack, and either the LLM returns llm_grounded_v1 or the rule interpreter returns "
        "rule_kb_v1. Dashboard tour recommendations elsewhere in SIGTS used heuristic scoring over a "
        "fixed activity catalogue (tags and season weights in browser storage), not an LLM. "
        "Administrative analytics anomaly flags used a simple z-score on daily sighting counts—"
        "classical statistics rather than neural networks."
    ),
]

TABLE_CAPTION = "Table 5.2: SIGTS tour-help artificial intelligence stack and algorithms"
TABLE_ROWS = [
    ("Generative AI (when API key set)", "GPT-4o or compatible chat model via OpenAI-compatible HTTP API", "Transformer language model (remote); temperature ~0.52; max tokens ~1024"),
    ("Grounding", "PostgreSQL + optional client catalogue snapshot", "Lexical RAG-lite: tokenise question, ILIKE retrieval, inject into system prompt"),
    ("Fallback / offline", "Node.js and browser rule interpreter", "Regex and keyword heuristics, template responses (rule_kb_v1)"),
    ("API surface", "REST POST /api/ai/chat, GET /api/ai/status", "JWT-authenticated; modes llm_grounded_v1, rule_kb_v1, rule_kb_v1_fallback"),
    ("Other “AI” features", "Dashboard recommendations; analytics anomalies", "Heuristic scoring; z-score on time series (not neural)"),
]


def para_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def set_times(p, size=12, bold=False, italic=False):
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_body_after(anchor, text):
    new_p = deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, anchor._parent)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    set_times(p)
    return p


def add_heading_after(anchor, title, level=4):
    new_p = deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, anchor._parent)
    p.style = f"Heading {level}"
    p.add_run(title)
    set_times(p, 12, bold=True)
    return p


def add_table_after(anchor, doc):
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technology"
    hdr[2].text = "Algorithm / method"
    for i, (a, b, c) in enumerate(TABLE_ROWS, 1):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
        tbl.rows[i].cells[2].text = c
    anchor._p.addnext(tbl._tbl)
    cap = add_body_after(anchor, TABLE_CAPTION)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_times(cap, 11, italic=True)
    return cap


def find_paragraph_index(doc, needle: str) -> int | None:
    n = needle.lower()
    for i, p in enumerate(doc.paragraphs):
        if n in para_text(p).lower():
            return i
    return None


def remove_range(doc, start: int, end: int):
    """Remove paragraphs [start, end) by clearing XML elements."""
    for i in range(end - 1, start - 1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)


def patch(doc: Document) -> bool:
    idx_523 = find_paragraph_index(doc, MARKER_523)
    idx_524 = find_paragraph_index(doc, MARKER_524)
    if idx_523 is None or idx_524 is None:
        return False

    idx_5231 = find_paragraph_index(doc, MARKER_5231)
    if idx_5231 is not None and idx_5231 < idx_524:
        remove_range(doc, idx_5231, idx_524)
        idx_524 = find_paragraph_index(doc, MARKER_524)
        if idx_524 is None:
            return False

    anchor = doc.paragraphs[idx_523]
    intro = doc.paragraphs[idx_523 + 1] if idx_523 + 1 < len(doc.paragraphs) else anchor
    if idx_523 + 1 < idx_524:
        intro._element.getparent().remove(intro._element)

    anchor = doc.paragraphs[idx_523]
    last = add_body_after(anchor, TOUR_HELP_INTRO)
    last = add_heading_after(last, MARKER_5231, level=4)
    for block in AI_PARAS:
        last = add_body_after(last, block)
    last = add_table_after(last, doc)
    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", default=DEFAULT_PATH)
    args = parser.parse_args()

    try:
        doc = Document(args.path)
    except Exception as e:
        print(f"Cannot open {args.path}: {e}", file=sys.stderr)
        print("Place your file at that path, or run: python scripts/generate_final_project_docx.py", file=sys.stderr)
        sys.exit(1)

    if not patch(doc):
        print(f"Could not find sections {MARKER_523!r} and {MARKER_524!r} in {args.path}", file=sys.stderr)
        sys.exit(1)

    doc.save(args.path)
    print(f"Patched AI section in {args.path}")


if __name__ == "__main__":
    main()
