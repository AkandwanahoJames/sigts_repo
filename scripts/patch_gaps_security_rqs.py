"""
Close two reviewer-identified gaps in the FINAL PROJECT report (excluding the
testing and ethics items, by request):

  1. Explicit Research Questions  -> new "1.3.3 Research Questions" in Chapter 1,
     inserted after the Specific Objectives and before "1.4 Research Significance".
  2. Security architecture depth  -> new "4.8 Security Architecture" at the end of
     Chapter 4, with a STRIDE threat model (Table 4.18) and an RBAC matrix
     (Table 4.19), grounded in the system's actual stack (JWT, bcrypt, optional
     TOTP MFA, RBAC, helmet, express-validator, rate limiting, audit logging).

Also registers the two new tables in the List of Tables and flags the document
so Word refreshes the table of contents on open.

Idempotent. Usage:
  python scripts/patch_gaps_security_rqs.py
  python scripts/patch_gaps_security_rqs.py --path "FINAL PROJECT_backup_revised.docx"
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEFAULT_PATH = Path(__file__).resolve().parents[1] / "FINAL PROJECT_backup_revised.docx"

RQ_HEADING = "1.3.3 Research Questions"
RQ_LEADIN = "In line with the specific objectives, the study sought to answer the following research questions:"
RQ_ITEMS = [
    "1. What system and user requirements are needed to develop a Smart Information Guide Tour System for Bwindi Impenetrable National Park?",
    "2. How can the Smart Information Guide Tour System be designed to meet the identified requirements of tourists, tour guides, and park management?",
    "3. How can the designed Smart Information Guide Tour System be implemented using appropriate web and geospatial technologies?",
    "4. To what extent does the implemented system satisfy the identified user requirements?",
]

SEC_HEADING = "4.8 Security Architecture"
SEC_INTRO = (
    "The Smart Information Guide Tour System (SIGTS) was designed using a "
    "defence-in-depth approach in which complementary security controls were "
    "layered across the client, the application programming interface (API), and "
    "the database tiers. Authentication was enforced using JSON Web Token (JWT) "
    "access and refresh tokens, with passwords hashed using bcrypt and an optional "
    "time-based one-time password (TOTP) providing multi-factor authentication. "
    "Authorisation followed role-based access control (RBAC), restricting each "
    "request to one of three roles, namely tourist, tour guide, and IT manager. "
    "At the transport layer, all traffic was secured using HTTPS/TLS. At the "
    "application layer, the API applied Helmet security headers, server-side input "
    "validation, parameterised SQL queries, and rate limiting on authentication and "
    "general endpoints. At the data layer, PostgreSQL constraints and regular "
    "backups preserved integrity, while the hosting platform encrypted data at "
    "rest. All privileged actions were recorded in an audit log to support "
    "accountability."
)
SEC_STRIDE_LEADIN = (
    "A STRIDE threat model was used to analyse the principal security threats to "
    "the system and the corresponding mitigations applied, as summarised in "
    "Table 4.18."
)
SEC_RBAC_LEADIN = (
    "Access to system capabilities was governed by the role-based access control "
    "matrix presented in Table 4.19, which maps each role to the actions it was "
    "permitted to perform."
)

STRIDE_CAPTION = "Table 4.18: STRIDE threat model for SIGTS"
STRIDE_HEADER = ["Threat category", "Example threat", "Mitigation in SIGTS"]
STRIDE_ROWS = [
    ["Spoofing", "Impersonating an IT manager or tour guide account",
     "JWT authentication, bcrypt password hashing, and optional TOTP multi-factor authentication"],
    ["Tampering", "Unauthorised alteration of wildlife content or tour records",
     "Server-side validation, parameterised queries, role-based write permissions, and audit logs"],
    ["Repudiation", "A user denying an action such as a content change",
     "Timestamped audit logs and activity records retained for accountability"],
    ["Information disclosure", "Exposure of tourist or staff personal data",
     "TLS in transit, encryption at rest, and role-restricted, least-privilege queries"],
    ["Denial of service", "Flooding the authentication or query endpoints",
     "Rate limiting, request size limits, and database connection pooling"],
    ["Elevation of privilege", "A tourist attempting to reach administrative routes",
     "RBAC authorisation middleware enforcing the principle of least privilege"],
]

RBAC_CAPTION = "Table 4.19: Role-based access control matrix"
RBAC_HEADER = ["Capability / resource", "Tourist", "Tour guide", "IT manager"]
RBAC_ROWS = [
    ["View tourism content and maps", "Yes", "Yes", "Yes"],
    ["Submit feedback and ratings", "Yes", "Yes", "Yes"],
    ["Report wildlife sightings", "Yes", "Yes", "Yes"],
    ["Manage tours and monitor guests", "No", "Yes", "Yes"],
    ["Generate tour reports", "No", "Yes", "Yes"],
    ["Manage and approve content (CMS)", "No", "No", "Yes"],
    ["Manage user accounts", "No", "No", "Yes"],
    ["Configure settings and monitor system health", "No", "No", "Yes"],
    ["View audit logs and manage backups", "No", "No", "Yes"],
]

LOT_AFTER = "Table 4.17: REST API endpoint catalogue for SIGTS"
LOT_NEW = [
    "Table 4.18: STRIDE threat model for SIGTS",
    "Table 4.19: Role-based access control matrix",
]


def style_exists(doc, name):
    return any(s.name == name for s in doc.styles)


def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


def make_para(doc, text, style):
    p = doc.add_paragraph(text)
    if style and style_exists(doc, style):
        p.style = doc.styles[style]
    return p


def set_cell(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def build_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    if style_exists(doc, "Table Grid"):
        table.style = doc.styles["Table Grid"]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, label in enumerate(header):
        set_cell(table.rows[0].cells[i], label, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, bold=(i == 0))
    return table


def insert_research_questions(doc):
    if any(p.text.strip() == RQ_HEADING for p in doc.paragraphs):
        return False
    anchor = find_para(
        doc,
        lambda p: p.style.name.startswith("Heading") and "research significance" in p.text.lower(),
    )
    if anchor is None:
        return False
    heading = make_para(doc, RQ_HEADING, "Heading 3")
    leadin = make_para(doc, RQ_LEADIN, "Body Text")
    items = [make_para(doc, t, "Body Text") for t in RQ_ITEMS]
    spacer = make_para(doc, "", "Normal")
    for el in [heading, leadin, *items, spacer]:
        anchor._p.addprevious(el._p)
    return True


def insert_security_section(doc):
    if any(p.text.strip() == SEC_HEADING for p in doc.paragraphs):
        return False
    anchor = find_para(
        doc,
        lambda p: p.style.name.startswith("Heading") and p.text.strip().upper().startswith("CHAPTER FIVE"),
    )
    if anchor is None:
        return False

    caption_style = "My Tables" if style_exists(doc, "My Tables") else "Normal"
    heading = make_para(doc, SEC_HEADING, "Heading 2")
    intro = make_para(doc, SEC_INTRO, "Body Text")
    stride_lead = make_para(doc, SEC_STRIDE_LEADIN, "Body Text")
    stride_cap = make_para(doc, STRIDE_CAPTION, caption_style)
    stride_tbl = build_table(doc, STRIDE_HEADER, STRIDE_ROWS)
    rbac_lead = make_para(doc, SEC_RBAC_LEADIN, "Body Text")
    rbac_cap = make_para(doc, RBAC_CAPTION, caption_style)
    rbac_tbl = build_table(doc, RBAC_HEADER, RBAC_ROWS)
    spacer = make_para(doc, "", "Normal")

    order = [heading._p, intro._p, stride_lead._p, stride_cap._p, stride_tbl._tbl,
             rbac_lead._p, rbac_cap._p, rbac_tbl._tbl, spacer._p]
    for el in order:
        anchor._p.addprevious(el)
    return True


def update_list_of_tables(doc):
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == LOT_AFTER and p.style.name == "Normal":
            anchor = p
            break
    if anchor is None:
        return False
    existing = {p.text.strip() for p in doc.paragraphs if p.style.name == "Normal"}
    # insert in order right after the 4.17 entry
    prev = anchor
    for entry in LOT_NEW:
        if entry in existing:
            continue
        new_p = doc.add_paragraph(entry)
        new_p.style = doc.styles["Normal"]
        new_p.paragraph_format.left_indent = anchor.paragraph_format.left_indent
        prev._p.addnext(new_p._p)
        prev = new_p
    return True


def enable_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", type=Path, default=DEFAULT_PATH)
    args = parser.parse_args()
    if not args.path.exists():
        print(f"File not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(args.path))
    rq = insert_research_questions(doc)
    sec = insert_security_section(doc)
    lot = update_list_of_tables(doc)
    enable_update_fields(doc)
    doc.save(str(args.path))

    print(f"Patched {args.path.name}:")
    print(f"  - research questions (1.3.3) inserted: {rq}")
    print(f"  - security architecture (4.8 + Tables 4.18, 4.19) inserted: {sec}")
    print(f"  - list of tables updated: {lot}; TOC set to refresh on open")


if __name__ == "__main__":
    main()
