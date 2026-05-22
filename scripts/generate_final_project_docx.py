"""
Generate SIGTS final project Word report — Makerere BIS format aligned to
GROUP 34 SNAPSKIN END OF YEAR FINAL PROJECT REPORT.pdf (structure, numbering, tense).
Original SIGTS content; not copied Snapskin prose.

Default output: FINAL PROJECT REVISED - 2.docx
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT_PATH = r"c:\Projects\SIGTS\FINAL PROJECT REVISED - 2.docx"
REPO_ROOT = Path(__file__).resolve().parents[1]
FIGURE_AI_PNG = REPO_ROOT / "docs" / "figures" / "figure-5-1-sigts-ai-architecture.png"
FIGURE_AI_SVG = REPO_ROOT / "docs" / "figures" / "figure-5-1-sigts-ai-architecture.svg"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_break(doc):
    doc.add_page_break()


def center_title(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size, bold)


def list_line(doc, text):
    """TOC / list entry with dot leader tab."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    parts = text.rsplit("\t", 1)
    r1 = p.add_run(parts[0])
    set_font(r1)
    if len(parts) > 1:
        p.add_run("\t")
        r2 = p.add_run(parts[1])
        set_font(r2)


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for fld_type, text in [
        ("begin", None),
        (None, ' TOC \\o "1-3" \\h \\z \\u '),
        ("separate", None),
        ("end", None),
    ]:
        if fld_type == "begin":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "begin")
            run._r.append(el)
        elif fld_type == "separate":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "separate")
            run._r.append(el)
        elif fld_type == "end":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "end")
            run._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)


def body(doc, text):
    for block in text.strip().split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_font(run)


def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, 11, italic=True)


def add_centered_figure(doc, image_path: Path, caption_text: str, width_inches: float = 6.4):
    """Embed architecture diagram when PNG exists (run generate_ai_architecture_figure.py first)."""
    if not image_path.is_file():
        body(
            doc,
            f"[Insert Figure: {image_path.name} — generate with: "
            "python scripts/generate_ai_architecture_figure.py]",
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption(doc, caption_text)


def chapter(doc, title):
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        set_font(run, 14, bold=True)


def section(doc, title, level=2):
    h = doc.add_heading(title, level=level)
    for run in h.runs:
        set_font(run, 12, bold=True)


def numbered_intro(doc, number_label, text):
    p = doc.add_paragraph()
    r0 = p.add_run(f"{number_label} ")
    set_font(r0, bold=True)
    r1 = p.add_run(text)
    set_font(r1)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15


# —— Front-matter lists (must match captions inside the report) ——
TOC_ENTRIES = [
    ("Declaration:", "i"),
    ("Approval:", "ii"),
    ("Dedication:", "iii"),
    ("Acknowledgement:", "iv"),
    ("List Of Tables:", "viii"),
    ("List Of Figures:", "ix"),
    ("List Of Acronyms:", "xi"),
    ("Abstract:", "xii"),
    ("Chapter One: Introduction", "1"),
    ("1.0 Introduction:", "1"),
    ("1.1 Background:", "1"),
    ("1.2 Statement Of The Problem:", "3"),
    ("1.3 Aims And Objectives", "3"),
    ("1.3.1 Main Objective", "3"),
    ("1.3.2 Specific Objectives:", "3"),
    ("1.4 Research Scope", "4"),
    ("1.4.1 Geographical Scope", "4"),
    ("1.4.2 Time Scope", "4"),
    ("1.4.3 Technical Scope", "4"),
    ("1.5 Significance Of Study", "4"),
    ("Chapter Two: Literature Review", "6"),
    ("2.0 Introduction:", "6"),
    ("2.1 Ecotourism And Visitor Information Systems", "6"),
    ("2.2 Mobile And Web Guides In Protected Areas", "7"),
    ("2.3 Related Applications And Comparative Review", "10"),
    ("2.4 Summary", "12"),
    ("Chapter Three: Research Methodology", "13"),
    ("3.0 Introduction:", "13"),
    ("3.1 Research Design", "13"),
    ("3.2 Research Methods", "14"),
    ("3.3 Target Population", "14"),
    ("3.4 Data Collection Methods", "15"),
    ("3.5 Data Analysis", "15"),
    ("3.6 System Design", "16"),
    ("3.7 System Implementation", "17"),
    ("3.8 System Security", "17"),
    ("3.9 System Testing And Validation", "18"),
    ("3.10 Summary", "18"),
    ("Chapter Four: System Analysis And Design", "20"),
    ("4.0 Introduction", "20"),
    ("4.1 System Study", "20"),
    ("4.2 System Analysis", "21"),
    ("4.3 Findings From Our Research", "21"),
    ("4.4 Software Specifications", "26"),
    ("4.5 Hardware Specification", "26"),
    ("4.6 User Requirements", "27"),
    ("4.7 Functional Requirements", "27"),
    ("4.8 Non-Functional Requirements", "28"),
    ("4.9 System Design", "28"),
    ("Chapter Five: System Implementation, Testing And Validation", "32"),
    ("5.0 Introduction", "32"),
    ("5.1 System Implementation", "32"),
    ("5.1.1 System Technologies", "32"),
    ("5.2 System Functionality", "33"),
    ("5.2.1 Authentication", "33"),
    ("5.2.2 Map Guidance And Sightings", "34"),
    ("5.2.3 Tour Help And Catalogues", "34"),
    ("5.2.3.1 Artificial Intelligence Architecture And Algorithms", "34"),
    ("5.2.4 Guide And IT Operations", "35"),
    ("5.3 System Testing", "35"),
    ("5.3.1 User Testing", "35"),
    ("5.4 System Validation", "36"),
    ("Chapter Six: Summary, Limitation, Conclusion And Recommendations", "37"),
    ("6.0 Introduction:", "37"),
    ("6.1 Achievements:", "37"),
    ("6.2 Challenges:", "38"),
    ("6.3 Limitation Of The Study:", "38"),
    ("6.4 Recommendations:", "38"),
    ("6.5 Conclusions:", "39"),
    ("References:", "40"),
    ("Appendix 1: Interview Guide For Field Respondents", "42"),
    ("Appendix 2: Questionnaire", "43"),
]

LIST_OF_TABLES = [
    "Table 2.1: Strengths and weaknesses of related park-guide applications",
    "Table 2.2: Related visitor-information applications",
    "Table 3.1: Summary of the techniques and tools used",
    "Table 4.1: Showing gender distribution of respondents",
    "Table 4.2: Source of information about the park before visit",
    "Table 4.3: Physical design for tourists",
    "Table 4.4: Physical design for guides",
    "Table 4.5: Physical design for IT managers",
    "Table 5.1: Shows the target users selected for user testing",
    "Table 5.2: SIGTS tour-help artificial intelligence stack and algorithms",
]

LIST_OF_FIGURES = [
    "Figure 5.1: SIGTS Tour Help artificial intelligence architecture",
    "Figure 2.1: Mobile device used for map-based park guidance",
    "Figure 3.1: Agile software development phases adopted for SIGTS",
    "Figure 4.1: Chart showing the percentages of both male and female respondents",
    "Figure 4.2: Chart showing age groups of respondents",
    "Figure 4.3: Different interfaces for different users",
    "Figure 4.4: The system architecture of the Smart Information Guide Tour System",
    "Figure 4.5: Context diagram (Level 0) for SIGTS",
    "Figure 4.6: Illustration of use case diagram for SIGTS",
    "Figure 4.7: Entity relationship diagram for SIGTS",
    "Figure 4.8: Login form for SIGTS",
    "Figure 4.9: Tourist dashboard for SIGTS",
    "Figure 4.10: Interactive map with park boundary and points of interest",
    "Figure 4.11: Guide operations dashboard for SIGTS",
    "Figure 4.12: IT manager admin dashboard with database user directory",
]

LIST_OF_ACRONYMS = [
    ("API", "Application Programming Interface"),
    ("APP", "Application"),
    ("BIS", "Bachelor of Information Systems"),
    ("CSS", "Cascading Style Sheets"),
    ("DFD", "Data Flow Diagram"),
    ("ERD", "Entity Relationship Diagram"),
    ("FK", "Foreign Key"),
    ("GPS", "Global Positioning System"),
    ("GUI", "Graphical User Interface"),
    ("HTML", "HyperText Markup Language"),
    ("IST", "Information Systems and Technology"),
    ("IT", "Information Technology"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("PK", "Primary Key"),
    ("POI", "Point of Interest"),
    ("PWA", "Progressive Web Application"),
    ("REST", "Representational State Transfer"),
    ("SDLC", "System Development Life Cycle"),
    ("SIGTS", "Smart Information Guide Tour System"),
    ("SQL", "Structured Query Language"),
    ("UWA", "Uganda Wildlife Authority"),
    ("WHO", "World Health Organization"),
]


def build(out_path: str | None = None):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # Title page
    center_title(doc, "SMART INFORMATION GUIDE TOUR SYSTEM (SIGTS)", 14, True)
    center_title(doc, "BWINDI IMPENETRABLE NATIONAL PARK")
    center_title(doc, "")
    center_title(doc, "BY")
    center_title(doc, "[PROJECT GROUP — INSERT NAMES AND REG. NUMBERS]")
    center_title(doc, "DEPARTMENT OF INFORMATION SYSTEMS")
    center_title(doc, "SCHOOL OF COMPUTING AND INFORMATICS TECHNOLOGY")
    center_title(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A Report submitted to the School of Computing and Informatics Technology\n"
        "for the study leading to a project in partial fulfilment of the\n"
        "requirements for the award of the Degree of Bachelor of Information Systems\n"
        "and Technology at Makerere University."
    )
    set_font(r)
    center_title(doc, "")
    center_title(doc, "Supervisor")
    center_title(doc, "[SUPERVISOR NAME AND CONTACT]")
    center_title(doc, "")
    center_title(doc, "[Month, Year]")
    add_break(doc)

    center_title(doc, "DECLARATION", 12, True)
    body(
        doc,
        "As members of the project group, we hereby declare that the information presented in "
        "this project report is our original work and that, to the best of our knowledge, no "
        "similar study has been conducted nor has the same report ever been presented to another "
        "university for the award of the Degree of Bachelor of Information Systems and Technology "
        "of Makerere University.",
    )
    doc.add_paragraph("Date: _________________________")
    add_break(doc)

    center_title(doc, "APPROVAL", 12, True)
    body(
        doc,
        "This is to certify that I have read, supervised, and approved the report for the project "
        "undertaken by the above group in partial fulfilment of the requirements for the award "
        "of the Degree of Bachelor of Information Systems and Technology of Makerere University.",
    )
    doc.add_paragraph("Signature: _________________________    Date: _________________________")
    add_break(doc)

    center_title(doc, "DEDICATION", 12, True)
    body(
        doc,
        "It is with genuine gratitude that we dedicate this work to our families, our supervisor, "
        "and the communities surrounding Bwindi Impenetrable National Park whose knowledge and "
        "hospitality informed the direction of the Smart Information Guide Tour System.",
    )
    add_break(doc)

    center_title(doc, "ACKNOWLEDGEMENT", 12, True)
    body(
        doc,
        "The completion of this undertaking could not have been possible without the participation "
        "and assistance of many people. We express our deep appreciation to our supervisor, the "
        "project coordinator, lecturers at the College of Computing and Information Sciences, "
        "Uganda Wildlife Authority stakeholders, field guides who assisted with data collection, "
        "and classmates who supported user testing of the prototype.",
    )
    add_break(doc)

    center_title(doc, "TABLE OF CONTENTS:", 12, True)
    for title, page in TOC_ENTRIES:
        dots = "." * max(4, 60 - len(title) - len(page))
        list_line(doc, f"{title} {dots}\t{page}")
    p = doc.add_paragraph()
    r = p.add_run(
        "(Update page numbers in Microsoft Word after final pagination: References → Update Table.)"
    )
    set_font(r, 10, italic=True)
    add_break(doc)

    center_title(doc, "LIST OF TABLES:", 12, True)
    for item in LIST_OF_TABLES:
        list_line(doc, item)
    add_break(doc)

    center_title(doc, "LIST OF FIGURES:", 12, True)
    for item in LIST_OF_FIGURES:
        list_line(doc, item)
    add_break(doc)

    center_title(doc, "LIST OF ACRONYMS", 12, True)
    for abbr, meaning in LIST_OF_ACRONYMS:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{abbr}:\t\t")
        set_font(r1, bold=True)
        r2 = p.add_run(meaning)
        set_font(r2)
    add_break(doc)

    center_title(doc, "ABSTRACT:", 12, True)
    body(
        doc,
        "Visitor numbers at Bwindi Impenetrable National Park continue to grow, yet printed "
        "brochures and informal briefings do not always provide timely, personalised guidance "
        "for tourists, guides, and administrative staff. This report presented the Smart "
        "Information Guide Tour System (SIGTS), a web-based application that combined species "
        "and cultural catalogues, map guidance, wildlife sighting records, guide workflows, "
        "and IT administrative analytics in a single platform backed by PostgreSQL and PostGIS.",
    )
    body(
        doc,
        "The system was implemented using HyperText Markup Language (HTML), Cascading Style "
        "Sheets (CSS), and JavaScript on the client; Express.js on Node.js for the application "
        "programming interface; and role-based access secured with JSON Web Tokens. Offline "
        "caching and synchronisation queues were incorporated so that brief loss of connectivity "
        "in the forest did not prevent reading cached content or deferring sightings until the "
        "device reconnected. User testing and validation confirmed that tourists, guides, and IT "
        "managers could complete essential tasks described in the requirements.",
    )
    body(
        doc,
        "In conclusion, SIGTS demonstrated that a maintainable, spatially aware information "
        "system could support park operations and visitor orientation when content was governed "
        "responsibly and infrastructure was monitored. Recommendations included pilot deployment "
        "with trained guides, scheduled content reviews, and further integration with official "
        "permit and ranger systems.",
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "Keywords: Application, Bwindi, Ecotourism, Map guidance, Park information system, "
        "SIGTS, Wildlife sightings"
    )
    set_font(r, italic=True)
    add_break(doc)

    # —— CHAPTER ONE ——
    chapter(doc, "CHAPTER ONE: INTRODUCTION")
    numbered_intro(
        doc,
        "1.0 Introduction:",
        "This chapter introduced the Smart Information Guide Tour System (SIGTS) developed for "
        "Bwindi Impenetrable National Park. It outlined the background to visitor information "
        "challenges, stated the problem addressed by the project, listed aims and objectives, "
        "defined the research scope, and explained the significance of the study.",
    )
    section(doc, "1.1 Background:")
    body(
        doc,
        "Bwindi Impenetrable National Park is a UNESCO World Heritage Site renowned for mountain "
        "gorillas and rich biodiversity. Tourists, guides, and park administrators required "
        "accurate, up-to-date information about species, safety, culture, routes, and operational "
        "status. Traditional paper-based materials were costly to revise and could not reflect "
        "recent sightings or administrative metrics. Digital guide systems used in other protected "
        "areas showed that map-centred, role-aware applications could improve preparation and "
        "on-site orientation when connectivity and content governance were planned.",
    )
    section(doc, "1.2 Statement Of The Problem:")
    body(
        doc,
        "Visitors often arrived without consolidated guidance on park rules, trail options, or "
        "cultural expectations. Guides lacked a single workspace for schedules, guest context, and "
        "sighting logs tied to authoritative species records. IT staff needed operational "
        "summaries—user activity, sightings volume, synchronisation queues—without manual "
        "database queries. The problem addressed by this project was therefore the absence of an "
        "integrated, maintainable information system tailored to Bwindi stakeholders.",
    )
    section(doc, "1.3 Aims And Objectives")
    section(doc, "1.3.1 Main Objective", 3)
    body(
        doc,
        "The main objective was to design, implement, test, and validate a Smart Information "
        "Guide Tour System that improved access to park information for tourists, guides, and IT "
        "managers at Bwindi Impenetrable National Park.",
    )
    section(doc, "1.3.2 Specific Objectives:", 3)
    body(
        doc,
        "The specific objectives were to: (i) review literature on ecotourism and park information "
        "systems; (ii) gather requirements from stakeholders; (iii) produce analysis and design "
        "models for SIGTS; (iv) implement a working prototype with map, catalogue, sightings, "
        "guide, and admin modules; (v) test and validate the system against requirements; and "
        "(vi) document findings, limitations, and recommendations.",
    )
    section(doc, "1.4 Research Scope")
    section(doc, "1.4.1 Geographical Scope", 3)
    body(doc, "The study focused on Bwindi Impenetrable National Park, Uganda, including sector gates and trail context represented in the system map.")
    section(doc, "1.4.2 Time Scope", 3)
    body(doc, "The project was conducted within the academic project period allocated for the Bachelor of Information Systems and Technology programme.")
    section(doc, "1.4.3 Technical Scope", 3)
    body(
        doc,
        "The deliverable was a web application with progressive-web features, a REST application "
        "programming interface, and a PostgreSQL database with spatial extensions. Native mobile "
        "store distribution and payment integrations were outside the scope of this study.",
    )
    section(doc, "1.5 Significance Of Study")
    body(
        doc,
        "The study benefited tourists through better pre-visit preparation; guides through operational "
        "tools; park management through analytics; developers through a documented reference "
        "implementation; government and conservation agencies through improved information "
        "channels; and future researchers through an account of requirements, design, and "
        "evaluation in a protected-area context.",
    )
    add_break(doc)

    # —— CHAPTER TWO ——
    chapter(doc, "CHAPTER TWO: LITERATURE REVIEW")
    numbered_intro(
        doc,
        "2.0 Introduction:",
        "This chapter reviewed published work on ecotourism information systems, mobile guides in "
        "protected areas, and technologies relevant to SIGTS.",
    )
    section(doc, "2.1 Ecotourism And Visitor Information Systems")
    body(
        doc,
        "Ecotourism emphasised low-impact travel and interpretation of natural heritage. Prior "
        "studies showed that timely information reduced risky behaviour and improved satisfaction "
        "when content was authoritative and culturally sensitive.",
    )
    section(doc, "2.2 Mobile And Web Guides In Protected Areas")
    body(
        doc,
        "Several commercial and research prototypes combined maps, species catalogues, and offline "
        "storage. Strengths included portability and search; weaknesses included outdated content "
        "and weak role separation between visitors and staff.",
    )
    section(doc, "2.3 Related Applications And Comparative Review")
    body(
        doc,
        "Comparable systems included national-park applications with trail maps, citizen-science "
        "platforms for wildlife observations, and museum-style interpretive guides. None combined "
        "the full set of Bwindi-specific geofencing, guide shift workflows, intranet staff tools, "
        "and IT predictive analytics targeted by SIGTS.",
    )
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["Application / approach", "Strengths", "Weaknesses"]):
        t.rows[0].cells[j].text = h
    data = [
        ("Generic park map applications", "Familiar map interface", "Limited local species depth"),
        ("Citizen-science observers", "Community sightings", "Weak official governance"),
        ("SIGTS (this project)", "Role-aware Bwindi stack", "Requires hosting and content care"),
    ]
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    caption(doc, "Table 2.1: Strengths and weaknesses of related park-guide applications")
    t2 = doc.add_table(rows=3, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Application"
    t2.rows[0].cells[1].text = "Relevance to SIGTS"
    t2.rows[1].cells[0].text = "National park interpretive apps"
    t2.rows[1].cells[1].text = "Map and content patterns"
    t2.rows[2].cells[0].text = "SIGTS"
    t2.rows[2].cells[1].text = "Integrated roles and live analytics"
    caption(doc, "Table 2.2: Related visitor-information applications")
    section(doc, "2.4 Summary")
    body(doc, "The literature supported development of an integrated, web-centred guide with spatial data and stakeholder-specific interfaces.")
    caption(doc, "Figure 2.1: Mobile device used for map-based park guidance")
    add_break(doc)

    # —— CHAPTER THREE ——
    chapter(doc, "CHAPTER THREE: RESEARCH METHODOLOGY")
    numbered_intro(
        doc,
        "3.0 Introduction:",
        "This chapter described how the study was planned, how data were collected and analysed, "
        "and how the system was designed, implemented, tested, and validated.",
    )
    section(doc, "3.1 Research Design")
    body(doc, "A pragmatic design combined qualitative interviews with structured questionnaires and iterative software development.")
    section(doc, "3.2 Research Methods")
    body(doc, "The team used document review, stakeholder interviews, observation of guide briefings, and prototype demonstrations.")
    section(doc, "3.3 Target Population")
    body(doc, "The target population included tourists, tour guides, park administrative staff, and information-technology officers supporting park systems.")
    section(doc, "3.4 Data Collection Methods")
    body(doc, "Questionnaires and interview guides were prepared, piloted, and administered to willing participants. Responses informed requirements in Chapter Four.")
    section(doc, "3.5 Data Analysis")
    body(doc, "Quantitative responses were summarised with descriptive statistics. Qualitative notes were coded by theme (navigation, safety, sightings, administration).")
    section(doc, "3.6 System Design")
    body(doc, "Design followed agile iterations: context diagrams, use cases, data models, and interface wireframes were refined each sprint.")
    caption(doc, "Figure 3.1: Agile software development phases adopted for SIGTS")
    section(doc, "3.7 System Implementation")
    body(doc, "Implementation is detailed in Chapter Five. Modular services were integrated behind a single application programming interface.")
    section(doc, "3.8 System Security")
    body(doc, "Security requirements included hashed passwords, token-based sessions, role authorisation, and production configuration checks.")
    section(doc, "3.9 System Testing And Validation")
    body(doc, "Testing combined unit checks, integration walkthroughs, and user rating sessions. Validation ensured inputs and outputs matched predefined formats.")
    section(doc, "3.10 Summary")
    body(doc, "The methodology linked stakeholder evidence to design, construction, and evaluation of SIGTS.")
    t3 = doc.add_table(rows=5, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Activity"
    t3.rows[0].cells[1].text = "Tool / technique"
    for i, row in enumerate(
        [
            ("Requirements gathering", "Interviews, questionnaires"),
            ("Design", "UML-style diagrams, wireframes"),
            ("Implementation", "Node.js, PostgreSQL, Leaflet"),
            ("Testing", "Manual scripts, Likert questionnaire"),
        ],
        1,
    ):
        t3.rows[i].cells[0].text = row[0]
        t3.rows[i].cells[1].text = row[1]
    caption(doc, "Table 3.1: Summary of the techniques and tools used")
    add_break(doc)

    # —— CHAPTER FOUR ——
    chapter(doc, "CHAPTER FOUR: SYSTEM ANALYSIS AND DESIGN")
    numbered_intro(
        doc,
        "4.0 Introduction",
        "This chapter presented analysis findings, requirements, and design artefacts for SIGTS.",
    )
    section(doc, "4.1 System Study")
    body(doc, "The existing process relied on briefings, signage, and informal communication. Bottlenecks included inconsistent species descriptions and delayed operational visibility.")
    section(doc, "4.2 System Analysis")
    body(doc, "Analysis produced actor lists, data flows, and priority requirements for tourists, guides, and IT managers.")
    section(doc, "4.3 Findings From Our Research")
    body(
        doc,
        "Respondents valued map-based orientation, species photographs, safety reminders, and cultural "
        "stories. Guides requested shift-friendly dashboards. Administrators requested user and "
        "sightings summaries without manual exports.",
    )
    t41 = doc.add_table(rows=3, cols=3)
    t41.style = "Table Grid"
    t41.rows[0].cells[0].text = "Gender"
    t41.rows[0].cells[1].text = "Frequency"
    t41.rows[0].cells[2].text = "Percentage"
    for ci, val in enumerate(["Male", "—", "—"]):
        t41.rows[1].cells[ci].text = val
    for ci, val in enumerate(["Female", "—", "—"]):
        t41.rows[2].cells[ci].text = val
    caption(doc, "Table 4.1: Showing gender distribution of respondents")
    caption(doc, "Figure 4.1: Chart showing the percentages of both male and female respondents")
    t42 = doc.add_table(rows=3, cols=2)
    t42.style = "Table Grid"
    t42.rows[0].cells[0].text = "Source"
    t42.rows[0].cells[1].text = "Count"
    t42.rows[1].cells[0].text = "Park website / staff"
    t42.rows[1].cells[1].text = "—"
    t42.rows[2].cells[0].text = "Friends / social media"
    t42.rows[2].cells[1].text = "—"
    caption(doc, "Table 4.2: Source of information about the park before visit")
    caption(doc, "Figure 4.2: Chart showing age groups of respondents")
    section(doc, "4.4 Software Specifications")
    body(doc, "Client: modern browser with JavaScript enabled. Server: Node.js 18+, PostgreSQL 14+ with PostGIS. Deployment: reverse proxy optional.")
    section(doc, "4.5 Hardware Specification")
    body(doc, "Server: multi-core host with solid-state storage. Clients: mid-range smartphones or laptops used by tourists and staff.")
    section(doc, "4.6 User Requirements")
    body(doc, "Users required intuitive navigation, readable typography outdoors, and clear role-specific menus.")
    section(doc, "4.7 Functional Requirements")
    body(
        doc,
        "Key functions included authentication, species catalogue, cultural content, interactive map, "
        "sighting submission, guide schedules, IT analytics, intranet announcements, inventory "
        "tracking, and an administrative user directory fed from the live database.",
    )
    section(doc, "4.8 Non-Functional Requirements")
    body(doc, "Non-functional requirements covered performance, security, offline tolerance, accessibility, and maintainability.")
    section(doc, "4.9 System Design")
    body(
        doc,
        "SIGTS adopted a three-tier architecture: presentation (web client), application (Express API), "
        "and data (PostgreSQL/PostGIS). Context, use case, sequence, and entity relationship diagrams "
        "were prepared. Screen designs covered tourist, guide, and IT dashboards.",
    )
    caption(doc, "Figure 4.3: Different interfaces for different users")
    caption(doc, "Figure 4.4: The system architecture of the Smart Information Guide Tour System")
    caption(doc, "Figure 4.5: Context diagram (Level 0) for SIGTS")
    caption(doc, "Figure 4.6: Illustration of use case diagram for SIGTS")
    caption(doc, "Figure 4.7: Entity relationship diagram for SIGTS")
    for cap in [
        "Table 4.3: Physical design for tourists",
        "Table 4.4: Physical design for guides",
        "Table 4.5: Physical design for IT managers",
    ]:
        caption(doc, cap)
    caption(doc, "Figure 4.8: Login form for SIGTS")
    caption(doc, "Figure 4.9: Tourist dashboard for SIGTS")
    caption(doc, "Figure 4.10: Interactive map with park boundary and points of interest")
    caption(doc, "Figure 4.11: Guide operations dashboard for SIGTS")
    caption(doc, "Figure 4.12: IT manager admin dashboard with database user directory")
    add_break(doc)

    # —— CHAPTER FIVE —— (PDF tense pattern)
    chapter(doc, "CHAPTER FIVE: SYSTEM IMPLEMENTATION, TESTING AND VALIDATION")
    numbered_intro(
        doc,
        "5.0 Introduction",
        "This chapter defines the implementation of the design to meet the requirements of the "
        "system as well as testing and validation of the system.",
    )
    section(doc, "5.1 System Implementation")
    section(doc, "5.1.1 System Technologies", 3)
    body(
        doc,
        "The Smart Information Guide Tour System was developed with the aid of standard web "
        "technologies. HyperText Markup Language (HTML) was used to specify the structure of "
        "documents retrieved through browser programs. Cascading Style Sheets (CSS) were used to "
        "control presentation. JavaScript (JS), an open-source client-side scripting language, was "
        "used to develop dynamic functionality. Its ability to display information in browsers "
        "enabled users to interact with the system.",
    )
    body(
        doc,
        "Leaflet was used for interactive mapping of gates, trails, and points of interest. The "
        "server layer was built on Express.js running on Node.js. PostgreSQL with PostGIS stored "
        "relational and spatial data. JSON Web Tokens secured authenticated sessions. A service "
        "worker cached selected assets for offline reading. Most of the software used in the "
        "application implementation was chosen because it was readily available and supported rapid "
        "development within the project timeline.",
    )
    section(doc, "5.2 System Functionality")
    section(doc, "5.2.1 Authentication", 3)
    body(
        doc,
        "The system allows tourists, guides, and IT managers to log in before accessing role-specific "
        "dashboards. The system requires registration and password recovery flows so that accounts "
        "could be provisioned without manual database edits. IT managers could deactivate accounts "
        "through the administrative directory.",
    )
    section(doc, "5.2.2 Map Guidance And Sightings", 3)
    body(
        doc,
        "The system allows users to view park boundaries, points of interest, routes, and recent "
        "wildlife sightings on the map module. Tourists could submit sightings that were stored in "
        "the central database.",
    )
    section(doc, "5.2.3 Tour Help And Catalogues", 3)
    body(
        doc,
        "The system allowed tourists to browse species and cultural catalogues and to receive tour-help "
        "responses through a conversational module (“Tour Help”) integrated in the progressive web "
        "application. Responses were not produced by a proprietary machine-learning model trained "
        "within the project; instead, the implementation combined (i) an optional hosted large "
        "language model (LLM) accessed through an OpenAI-compatible application programming interface, "
        "and (ii) a rule-based knowledge interpreter used when no API key was configured, when the "
        "network or provider failed, or when questions were filtered as clearly off-topic.",
    )
    section(doc, "5.2.3.1 Artificial Intelligence Architecture And Algorithms", 4)
    body(
        doc,
        "Figure 5.1 illustrates the end-to-end architecture of the Tour Help module, showing how "
        "the progressive web client, authenticated API layer, lexical retrieval service, optional "
        "hosted language model, and rule-based fallback interact. Table 5.2 summarises the same stack "
        "in tabular form.",
    )
    add_centered_figure(
        doc,
        FIGURE_AI_PNG,
        "Figure 5.1: SIGTS Tour Help artificial intelligence architecture",
    )
    body(
        doc,
        "The primary generative path, when enabled, used "
        "the Chat Completions protocol (POST /v1/chat/completions) with default model GPT-4o, "
        "configurable through environment variables (SIGTS_CHAT_OPENAI_API_KEY, SIGTS_CHAT_MODEL, "
        "SIGTS_CHAT_TEMPERATURE approximately 0.52, SIGTS_CHAT_MAX_TOKENS up to 1024). GPT-4o is "
        "a commercial transformer-based decoder-only language model; SIGTS invoked it as a remote "
        "service and did not implement neural training or inference locally.",
    )
    body(
        doc,
        "Before each LLM call, the backend performed retrieval-augmented generation in a lightweight "
        "form (RAG-lite). The visitor question was tokenised (lowercase words of at least three "
        "characters), converted to SQL ILIKE patterns, and matched against PostgreSQL tables "
        "including parks, FAQs, safety tips, destination information, animals, wildlife tour themes, "
        "cultural narratives, locations, tour routes, and recent sightings. A core Bwindi briefing "
        "block was always appended so the model retained baseline park facts even when lexical "
        "matches were weak. The client also supplied a sanitised catalogue snapshot (themes, animals, "
        "map points, stories, FAQs, safety species) assembled in the browser. Retrieved text was "
        "injected into a system prompt that instructed the model to stay on Bwindi visitor topics, "
        "paraphrase grounded facts, and decline to invent permit prices or citations. This approach "
        "did not use vector embeddings or a semantic search index; matching was keyword-based only.",
    )
    body(
        doc,
        "The fallback algorithm (modes labelled rule_kb_v1 and rule_kb_v1_fallback) was a "
        "deterministic expert-system style pipeline: regular-expression intent detection for themes "
        "such as gorillas, safety, weather, culture, and map routing; scope checks for Bwindi and "
        "nature-tourism vocabulary; substring matching against the on-device species catalogue; and "
        "template paragraphs for greetings and general redirection. The same rule order was "
        "mirrored in JavaScript for offline use when POST /api/ai/chat was unavailable. Eligible "
        "questions were logged in the ai_query_logs table with response time and language metadata. "
        "Configuration could be inspected at GET /api/ai/status (llm_configured, model, provider, "
        "grounding tables).",
    )
    body(
        doc,
        "The request flow was as follows: an authenticated tourist submitted a question; "
        "optional GPS coordinates resolve to the nearest map point of interest, retrieval builds the "
        "knowledge pack, and either the LLM returns llm_grounded_v1 or the rule interpreter returns "
        "rule_kb_v1. Dashboard tour recommendations elsewhere in SIGTS used heuristic scoring over a "
        "fixed activity catalogue (tags and season weights in browser storage), not an LLM. "
        "Administrative analytics anomaly flags used a simple z-score on daily sighting counts—"
        "classical statistics rather than neural networks.",
    )
    table_ai = doc.add_table(rows=6, cols=3)
    table_ai.style = "Table Grid"
    ai_hdr = table_ai.rows[0].cells
    ai_hdr[0].text = "Layer"
    ai_hdr[1].text = "Technology"
    ai_hdr[2].text = "Algorithm / method"
    ai_rows = [
        (
            "Generative AI (when API key set)",
            "GPT-4o or compatible chat model via OpenAI-compatible HTTP API",
            "Transformer language model (remote); temperature ~0.52; max tokens ~1024",
        ),
        (
            "Grounding",
            "PostgreSQL + optional client catalogue snapshot",
            "Lexical RAG-lite: tokenise question, ILIKE retrieval, inject into system prompt",
        ),
        (
            "Fallback / offline",
            "Node.js and browser rule interpreter",
            "Regex and keyword heuristics, template responses (rule_kb_v1)",
        ),
        (
            "API surface",
            "REST POST /api/ai/chat, GET /api/ai/status",
            "JWT-authenticated; modes llm_grounded_v1, rule_kb_v1, rule_kb_v1_fallback",
        ),
        (
            "Other “AI” features",
            "Dashboard recommendations; analytics anomalies",
            "Heuristic scoring; z-score on time series (not neural)",
        ),
    ]
    for i, (layer, tech, algo) in enumerate(ai_rows, 1):
        table_ai.rows[i].cells[0].text = layer
        table_ai.rows[i].cells[1].text = tech
        table_ai.rows[i].cells[2].text = algo
    caption(doc, "Table 5.2: SIGTS tour-help artificial intelligence stack and algorithms")
    section(doc, "5.2.4 Guide And IT Operations", 3)
    body(
        doc,
        "The system enables guides to access schedules and record sightings. IT managers could view "
        "predictive analytics, operational snapshots, and a graphical database user directory that "
        "loaded all accounts from the server.",
    )
    section(doc, "5.3 System Testing")
    body(
        doc,
        "The following testing strategies were deployed. System testing was done after the system "
        "was coded. Individual units or components of the system were checked to ensure they were "
        "fully functional before integration.",
    )
    section(doc, "5.3.1 User Testing", 3)
    body(
        doc,
        "A group of target users was selected to examine the system functionality and gave feedback "
        "as shown in Table 5.1 below.",
    )
    table5 = doc.add_table(rows=8, cols=7)
    table5.style = "Table Grid"
    hdr = table5.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "User testing question"
    for col, label in zip(range(2, 7), ["1", "2", "3", "4", "5"]):
        hdr[col].text = label
    questions = [
        "The system enables tourists to locate species and park information.",
        "The system enables tourists to use the map for gates, sectors, and points of interest.",
        "The system enables tourists to submit or review wildlife sightings.",
        "The system enables guides to access schedules and operational tools.",
        "The system enables IT managers to view analytics and database user summaries.",
        "The system enables tour-help responses relevant to Bwindi safety themes.",
        "The system enables use when the network connection is interrupted briefly.",
    ]
    for idx, q in enumerate(questions, 1):
        table5.rows[idx].cells[0].text = str(idx)
        table5.rows[idx].cells[1].text = q
    caption(doc, "Table 5.1: Shows the target users selected for user testing")
    section(doc, "5.4 System Validation")
    body(
        doc,
        "System validation is concerned with ensuring that data entered into the systems meets "
        "predefined formats with defined input criteria. It was done to ensure that the data entered "
        "and retrieved was valid.",
    )
    add_break(doc)

    # —— CHAPTER SIX —— (PDF headings)
    chapter(doc, "CHAPTER SIX: SUMMARY, LIMITATION, CONCLUSION AND RECOMMENDATIONS")
    numbered_intro(
        doc,
        "6.0 Introduction:",
        "This chapter provides a summary, limitation, conclusion, and recommendations of the "
        "Smart Information Guide Tour System.",
    )
    section(doc, "6.1 Achievements:")
    body(
        doc,
        "The Smart Information Guide Tour System ensured improved access to park information, "
        "map-based orientation, guide workflows, and administrative visibility for IT managers. "
        "Offline caching and synchronisation were implemented for field use.",
    )
    section(doc, "6.2 Challenges:")
    body(
        doc,
        "The SIGTS application was intended to be used on smartphones and browsers when there "
        "was internet access. As such, it required an internet connection for authentication and "
        "analytics. The system could not fully function if the server stopped running. The server "
        "was required to run when the system was loading administrative data.",
    )
    section(doc, "6.3 Limitation Of The Study:")
    body(
        doc,
        "The project needed a variety of resources, including funds to facilitate data collection "
        "and computer systems for developing the application, which were not always easy to obtain. "
        "It was difficult to obtain extended access to all field respondents during peak trekking "
        "seasons. Some specification items were documented as future work rather than completed "
        "within the project window.",
    )
    section(doc, "6.4 Recommendations:")
    body(
        doc,
        "SIGTS is highly recommended for pilot use by tourists, guides, and park administrators "
        "because its aim is to reduce the information gap between visitors and park authorities. "
        "It is also important to train and encourage users to become accustomed to the application. "
        "This being a new application, some users may need time to understand how the system functions.",
    )
    section(doc, "6.5 Conclusions:")
    body(
        doc,
        "The core reason for the establishment of the Smart Information Guide Tour System was to "
        "improve visitor orientation and operational visibility at Bwindi Impenetrable National Park. "
        "Therefore, the technologies used should support the core objective of the system if it is "
        "to remain relevant to the users. A lot still needs to be done in the IT department to make "
        "available technology effective. The application administration has to keep updating hardware "
        "and software requirements through its panel. Our research therefore acknowledges the fact "
        "that the system was delivered as a working prototype suitable for academic review and pilot "
        "demonstration.",
    )
    add_break(doc)

    center_title(doc, "REFERENCES", 12, True)
    body(
        doc,
        "[1] Uganda Wildlife Authority. Tourism and conservation publications.\n\n"
        "[2] Fielding, R. T. Architectural Styles and the Design of Network-based Software Architectures.\n\n"
        "[3] PostGIS Project. Spatial database extensions for PostgreSQL.\n\n"
        "[4] Additional peer-reviewed and grey literature cited during the project.",
    )
    add_break(doc)
    center_title(doc, "Appendix 1: Interview Guide For Field Respondents", 12, True)
    body(doc, "Insert interview guide used during data collection.")
    add_break(doc)
    center_title(doc, "Appendix 2: Questionnaire", 12, True)
    body(doc, "Insert questionnaire instrument.")

    target = out_path or OUT_PATH
    doc.save(target)
    return target


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate SIGTS final project Word report.")
    parser.add_argument(
        "--out",
        default=OUT_PATH,
        help="Output .docx path (default: FINAL PROJECT REVISED - 2.docx)",
    )
    args = parser.parse_args()
    print("Wrote", build(args.out))
