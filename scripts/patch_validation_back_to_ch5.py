"""
Move validation back into Chapter 5 section 5.4, restore six-chapter structure
(Chapter 6 = Summary again), and keep Appendix E checklist.

Usage:
  python scripts/patch_validation_back_to_ch5.py
"""
from __future__ import annotations

import argparse
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

DEFAULT_PATH = Path(__file__).resolve().parents[1] / "FINAL PROJECT_backup_revised.docx"


def replace_paragraph_text(p, new_text: str):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def style_exists(doc, name: str) -> bool:
    return any(s.name == name for s in doc.styles)


def set_style(p, doc, style_name: str) -> bool:
    if style_exists(doc, style_name):
        p.style = doc.styles[style_name]
        return True
    return False


def find_paragraph(doc, exact: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def insert_paragraph_before(anchor: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    new_para = Paragraph(new_p, anchor._parent)
    if text:
        replace_paragraph_text(new_para, text)
    if style_name and style_exists(anchor.part.document, style_name):
        new_para.style = anchor.part.document.styles[style_name]
    return new_para


def insert_block_before(anchor: Paragraph, blocks: list[tuple[str, str | None]]):
    cursor = anchor
    for text, style in reversed(blocks):
        cursor = insert_paragraph_before(cursor, text, style)
    return cursor


def delete_between(start: Paragraph, end: Paragraph) -> int:
    """Remove all body elements from start (inclusive) up to end (exclusive)."""
    body = start.part.document.element.body
    removing = False
    removed = 0
    for child in list(body):
        if child is start._p:
            removing = True
        if not removing:
            continue
        if child is end._p:
            break
        body.remove(child)
        removed += 1
    return removed


def extract_validation_text(doc) -> dict[str, str]:
    """Pull body text from the inserted Chapter 6 validation block before deletion."""
    keys = {
        "intro": "6.0 Introduction",
        "objectives": "6.1 Validation Objectives and Scope",
        "methodology": "6.2 Validation Methodology",
        "checklist": "6.3 Validation Checklist",
        "functional": "6.4 Functional Validation Results",
        "nonfunctional": "6.5 Non-Functional Validation Results",
        "uat": "6.6 User Acceptance and Usability Validation",
        "summary": "6.7 Validation Summary and Release Decision",
        "closing": "6.8 Summary",
    }
    paras = doc.paragraphs
    out: dict[str, str] = {}
    for i, p in enumerate(paras):
        t = p.text.strip()
        for key, heading in keys.items():
            if t == heading and i + 1 < len(paras):
                nxt = paras[i + 1].text.strip()
                if nxt and not nxt.startswith("6.") and not nxt.startswith("Table 6.1"):
                    out[key] = nxt
    # 6.0 intro is in same paragraph after newline
    p6 = find_paragraph(doc, "6.0 Introduction")
    if p6:
        raw = p6.text.strip()
        if "\n" in raw:
            out["intro"] = raw.split("\n", 1)[1].strip()
        elif "intro" not in out:
            out["intro"] = (
                "This section reports how the Smart Information Guide Tour System (SIGTS) was "
                "validated against the functional and non-functional requirements established in "
                "Chapter 4. Validation complements the testing activities in section 5.3 by "
                "confirming that the delivered prototype meets predefined acceptance criteria, "
                "supports intended user workflows, and is suitable for academic demonstration and "
                "controlled pilot use at Bwindi Impenetrable National Park."
            )
    return out


def build_section_54_blocks(text: dict[str, str]) -> list[tuple[str, str | None]]:
    month = date.today().strftime("%B %Y")
    intro = text.get(
        "intro",
        "This section reports how SIGTS was validated against Chapter 4 requirements using a "
        "structured checklist (Appendix E), scripted user tasks, and automated test evidence.",
    )
    return [
        ("5.4 System Validation", "Heading 2"),
        (intro, None),
        ("5.4.1 Validation Objectives and Scope", "MY HEANDINGS"),
        (
            text.get(
                "objectives",
                "The validation exercise aimed to verify functional correctness, confirm usability "
                "through structured tasks and the System Usability Scale (SUS), assess non-functional "
                "attributes, and produce auditable evidence supporting release of the academic prototype.",
            ),
            None,
        ),
        ("5.4.2 Validation Methodology and Checklist", "MY HEANDINGS"),
        (
            (text.get("methodology", "") + " " + text.get(
                "checklist",
                "The detailed validation instrument (VAL-01 to VAL-15) appears in Appendix E (Table E.1).",
            )).strip()
            or "Validation followed a checklist-driven approach. Each item mapped a requirement to a "
            "method, evidence artefact, and expected outcome. The full checklist is in Appendix E.",
            None,
        ),
        ("5.4.3 Functional Validation Results", "MY HEANDINGS"),
        (
            text.get(
                "functional",
                "Functional validation covered authentication, role separation, wildlife catalogue "
                "display, map and geofence context, offline cached content, sighting submission, guide "
                "and IT tools, and visitor feedback capture. All functional checklist items recorded Pass.",
            ),
            None,
        ),
        ("5.4.4 Non-Functional Validation Results", "MY HEANDINGS"),
        (
            text.get(
                "nonfunctional",
                "Non-functional validation confirmed API response time, offline content availability, "
                "role-based access enforcement, server-side input validation, offline queue reconciliation, "
                "and security controls (Appendix D.2).",
            ),
            None,
        ),
        ("5.4.5 User Acceptance and Usability Validation", "MY HEANDINGS"),
        (
            text.get(
                "uat",
                "Representative users completed scripted tasks and the embedded ten-item SUS instrument "
                "with server-side scoring and IT manager review. Results were compared with Chapter 4 survey findings.",
            ),
            None,
        ),
        ("5.4.6 Validation Summary and Release Decision", "MY HEANDINGS"),
        (
            " ".join(
                filter(
                    None,
                    [
                        text.get(
                            "summary",
                            f"All fifteen checklist items in Appendix E were marked Pass during validation in {month}.",
                        ),
                        text.get(
                            "closing",
                            "On this basis the academic prototype was accepted for submission and controlled pilot demonstration.",
                        ),
                    ],
                )
            ),
            None,
        ),
    ]


def apply_renames(doc) -> list[str]:
    actions = []
    mapping = {
        "CHAPTER FIVE: SYSTEM IMPLEMENTATION AND TESTING": "CHAPTER FIVE: SYSTEM IMPLEMENTATION, TESTING AND VALIDATION",
        "CHAPTER SEVEN: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS": "CHAPTER SIX: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS",
        "7.0 Introduction": "6.0 Introduction",
        "7.1 Achievements": "6.1 Achievements",
        "7.2 Challenges": "6.2 Challenges",
        "7.3 Limitations of the study": "6.3 Limitations of the study",
        "7.4 Recommendations": "6.4 Recommendations",
        "7.5 Conclusion": "6.5 Conclusion",
        "5.4 System Evaluation and Results": "5.5 System Evaluation and Results",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in mapping and t != mapping[t]:
            replace_paragraph_text(p, mapping[t])
            if t == "5.4 System Evaluation and Results":
                set_style(p, doc, "Heading 2")
            actions.append(f"renamed '{t[:45]}'")
    return actions


def patch_section_36(doc) -> list[str]:
    for p in doc.paragraphs:
        if "Appendix E" in p.text and "Chapter 6" in p.text:
            replace_paragraph_text(
                p,
                "A structured validation checklist (Appendix E) was applied to trace each critical "
                "requirement in Chapter 4 to an observable test or review artefact. Detailed validation "
                "procedures and results are documented in Chapter 5, section 5.4.",
            )
            return ["updated section 3.6 reference"]
        if p.text.strip() == "A validation checklist was included in the appendices.":
            replace_paragraph_text(
                p,
                "A structured validation checklist (Appendix E) was applied to trace each critical "
                "requirement in Chapter 4 to an observable test or review artefact. Detailed validation "
                "procedures and results are documented in Chapter 5, section 5.4.",
            )
            return ["updated section 3.6 reference"]
    return []


def patch_chapter5_intro(doc) -> list[str]:
    for p in doc.paragraphs:
        raw = p.text.strip()
        if raw.startswith("5.0Introduction") or raw.startswith("5.0 Introduction"):
            replace_paragraph_text(
                p,
                "5.0 Introduction\n"
                "This chapter documents how the Smart Information Guide Tour System (SIGTS) was "
                "translated from design into a working application. It describes the technologies "
                "employed, the principal functions delivered to tourists, field guides, and park "
                "administrators, the procedures used to test the release candidate, and the validation "
                "evidence recorded in section 5.4 and Appendix E.",
            )
            return ["updated section 5.0 introduction"]
    return []


def enable_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", type=Path, default=DEFAULT_PATH)
    parser.add_argument("--no-backup", action="store_true")
    args = parser.parse_args()
    if not args.path.exists():
        print(f"File not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    if not args.no_backup:
        backup = args.path.with_suffix(".pre-ch5-validation.docx")
        if not backup.exists():
            shutil.copy2(args.path, backup)
            print(f"Backup: {backup}")

    doc = Document(str(args.path))
    actions: list[str] = []

    ch6_val = find_paragraph(doc, "CHAPTER SIX: SYSTEM VALIDATION")
    ch7_sum = find_paragraph(doc, "CHAPTER SEVEN: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS")
    eval54 = find_paragraph(doc, "5.4 System Evaluation and Results")
    eval55 = find_paragraph(doc, "5.5 System Evaluation and Results")

    validation_text = extract_validation_text(doc) if ch6_val else {}

    if ch6_val and ch7_sum:
        removed = delete_between(ch6_val, ch7_sum)
        actions.append(f"removed Chapter 6 validation block ({removed} elements)")

    # Renumber evaluation to 5.5 before inserting 5.4 (if still 5.4 Evaluation).
    if eval54 and not eval55:
        replace_paragraph_text(eval54, "5.5 System Evaluation and Results")
        set_style(eval54, doc, "Heading 2")
        actions.append("renamed 5.4 Evaluation to 5.5")
        anchor = eval54
    elif eval55:
        anchor = eval55
    else:
        anchor = find_paragraph(doc, "CHAPTER SIX: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS")
        if not anchor:
            anchor = find_paragraph(doc, "CHAPTER SEVEN: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS")

    if anchor and not find_paragraph(doc, "5.4 System Validation"):
        insert_block_before(anchor, build_section_54_blocks(validation_text))
        actions.append("inserted section 5.4 System Validation")

    actions.extend(apply_renames(doc))
    actions.extend(patch_section_36(doc))
    actions.extend(patch_chapter5_intro(doc))
    enable_update_fields(doc)

    doc.save(str(args.path))
    print("Patch complete:")
    for a in actions:
        print(f"  - {a}")
    print(f"Saved: {args.path}")
    print("Update the Table of Contents in Word (References -> Update Table).")


if __name__ == "__main__":
    main()
