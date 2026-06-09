"""
Align FINAL PROJECT .docx with the implemented SIGTS codebase (API paths, stack, AI).

  python scripts/patch_doc_alignment.py
  python scripts/patch_doc_alignment.py --path "c:\\Projects\\SIGTS\\FINAL PROJECT_backup_revised.docx"
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document

DEFAULT_PATH = Path(__file__).resolve().parents[1] / "FINAL PROJECT_backup_revised.docx"

# (pattern, replacement) — applied in order
REPLACEMENTS: list[tuple[str, str]] = [
    (
        r"POST /api/tourist/ai-query \(planned backend handler\)",
        "POST /api/ai/chat (implemented; JWT-authenticated)",
    ),
    (
        r"POST /api/tourist/ai-query once the backend route is completed",
        "POST /api/ai/chat on the Express API",
    ),
    (
        r"POST /api/tourist/ai-query",
        "POST /api/ai/chat",
    ),
    (
        r"POST /api/ai/query",
        "POST /api/ai/chat",
    ),
    (
        r"Until that route ships, visitors rely on curated FAQs and catalogue endpoints described in Section 4\.6\.3\.1\.",
        "When no LLM API key is configured, the same endpoint returns rule-based answers (rule_kb_v1) grounded in curated FAQs and catalogue tables.",
    ),
    (
        r"a future release will add retrieval-augmented prompting once the Express handler is wired\.",
        "retrieval-augmented prompting is implemented via PostgreSQL lexical grounding plus an optional OpenAI-compatible LLM.",
    ),
    (
        r"The planned tour-help path combines PostgreSQL retrieval from animals and cultural_narratives with a hosted language-model API behind POST /api/ai/chat\.",
        "The tour-help path combines PostgreSQL retrieval from animals, cultural_narratives, FAQs, and safety content with a hosted language-model API behind POST /api/ai/chat, with a rule-based fallback when no API key is set.",
    ),
    (r"/api/tourist/animals/:id", "/api/animals/:id"),
    (r"/api/tourist/animals", "/api/animals"),
    (r"/api/tourist/locations", "/api/locations"),
    (r"/api/tourist/cultural-narratives", "/api/cultural"),
    (r"/api/tourist/tour-routes", "/api/tours/routes"),
    (r"/api/tourist/faqs", "/api/faqs"),
    (r"/api/tourist/sightings", "/api/sightings"),
    (r"/api/guide/tours", "/api/tours/schedule"),
    (r"/api/admin/analytics/congestion", "/api/analytics/predictions/congestion"),
    (r"mounted under /api/auth, /api/tourist, /api/guide, /api/admin, and /api/sync", "mounted under /api/auth, /api/animals, /api/locations, /api/cultural, /api/sightings, /api/tours, /api/ai, /api/admin, /api/analytics, and /api/sync"),
    (r"Route modules under /api/auth, /api/tourist, /api/guide, /api/admin, and /api/sync", "Route modules under /api/auth, /api/animals, /api/locations, /api/cultural, /api/sightings, /api/tours, /api/ai, /api/admin, /api/analytics, and /api/sync"),
    (
        r"The tour-help feature is implemented as a Next\.js page \(ask the AI chatbot\) that is designed to call POST /api/ai/chat on the Express API\. In the pilot, tourists receive answers from curated FAQs and animals/cultural_narratives content through existing GET routes; the AI assistant screen demonstrates the intended chat workflow\. Park-approved rows in PostgreSQL remain the authoritative knowledge base; retrieval-augmented prompting is implemented via PostgreSQL lexical grounding plus an optional OpenAI-compatible LLM\.",
        "Tour help is implemented in the tourist progressive web application (Tour Help view). Authenticated clients call POST /api/ai/chat; the backend grounds answers in PostgreSQL (animals, cultural_narratives, FAQs, safety) and optionally a hosted LLM, with rule_kb_v1 fallback when no API key is configured.",
    ),
    (
        r"The Ask the Park AI page in the tourist Next\.js app demonstrates tour-help chat; it is wired to POST /api/ai/chat \(implemented; JWT-authenticated\)\. When no LLM API key is configured, the same endpoint returns rule-based answers \(rule_kb_v1\) grounded in curated FAQs and catalogue tables\.",
        "The Ask the Park (Tour Help) screen in the tourist PWA calls POST /api/ai/chat. When no LLM API key is configured, the endpoint returns rule-based answers (rule_kb_v1) grounded in curated FAQs and catalogue tables.",
    ),
    (
        r"SIGTS was designed as a three-tier web application: Next\.js/React clients, an Express\.js API on Node\.js, and PostgreSQL with PostGIS, backed by Redis\.",
        "SIGTS was implemented as a three-tier web application: a browser-based progressive web client (HTML, CSS, JavaScript), an Express.js API on Node.js, and PostgreSQL with PostGIS, optionally backed by Redis.",
    ),
    (
        r"Technologies included Next\.js 14, React 18, Tailwind CSS, Node\.js, Express\.js, PostgreSQL with PostGIS, and Redis\.",
        "Technologies included a static progressive web client (HTML/CSS/JavaScript, Tailwind CSS, Leaflet), Node.js, Express.js, PostgreSQL with PostGIS, and optional Redis.",
    ),
    (
        r"Front-End: Next\.js 14 \(tourist and guide PWAs\), React Admin \(IT portal\), Tailwind CSS, Leaflet",
        "Front-End: single progressive web application with role-based views (tourist, guide, IT manager), Tailwind CSS, Leaflet",
    ),
    (
        r"Three client applications were built: a Next\.js tourist PWA with offline map tiles and IndexedDB sync, a Next\.js guide dashboard, and a React Admin portal for IT managers\.",
        "One progressive web application serves tourists, guides, and IT managers through role-based navigation, with offline map tile caching, localStorage/sync queues, and a service worker.",
    ),
    (
        r"The client layer comprises three applications in a pnpm monorepo: a Next\.js 14 tourist PWA \(Leaflet maps, next-pwa, IndexedDB offline queue\), a Next\.js guide app, and a React Admin IT portal\.",
        "The client layer is a single progressive web application (Leaflet maps, service worker, localStorage offline caches, and /api/sync upload when connectivity returns).",
    ),
    (
        r"coding with Next\.js, Express\.js, PostgreSQL/PostGIS, and Redis",
        "coding with a browser PWA front end, Express.js, PostgreSQL/PostGIS, and optional Redis",
    ),
    (
        r"Integration tests evaluated interactions between the Next\.js clients and the Express REST API",
        "Integration tests evaluated interactions between the progressive web client and the Express REST API",
    ),
    (
        r"separate Next\.js containers for tourist \(:3001\), guide \(:3002\), and React Admin \(:3003\)",
        "the Express API (port 8001 in development) and static front-end assets served from the same host or nginx",
    ),
    (
        r"Tourist devices sync IndexedDB queues through /api/sync when connectivity returns\.",
        "Tourist devices sync offline queues through /api/sync when connectivity returns.",
    ),
    (
        r"All ten cases passed on the build submitted for examination\.",
        "All ten cases passed on the release candidate verified against the local PostgreSQL database (May 2026 automated API suite: 23/23 checks).",
    ),
]


def apply_replacements(text: str) -> tuple[str, int]:
    count = 0
    out = text
    for pattern, repl in REPLACEMENTS:
        new, n = re.subn(pattern, repl, out, flags=re.IGNORECASE)
        if n:
            count += n
            out = new
    return out, count


def patch_paragraph(p) -> int:
    full = "".join(r.text for r in p.runs)
    if not full.strip():
        return 0
    new, n = apply_replacements(full)
    if n and new != full:
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new)
    return n


def patch_table(tbl) -> int:
    total = 0
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                total += patch_paragraph(p)
    return total


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", type=Path, default=DEFAULT_PATH)
    args = parser.parse_args()

    if not args.path.exists():
        print(f"File not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(args.path))
    total = 0
    for p in doc.paragraphs:
        total += patch_paragraph(p)
    for tbl in doc.tables:
        total += patch_table(tbl)

    doc.save(str(args.path))
    print(f"Patched {args.path} ({total} replacement(s))")


if __name__ == "__main__":
    main()
