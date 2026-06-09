"""
Fix Chapter 4 numbering/TOC issues introduced and pre-existing around the new
Security Architecture section in the FINAL PROJECT report.

Changes (idempotent):
  1. Resolve the duplicate "4.8": rename the new "4.8 Security Architecture" to
     "4.9 Security Architecture" (it follows the pre-existing "4.8 Data Inputs").
  2. Make "4.8 Data Inputs" appear in the table of contents by converting it from
     the custom "MY HEANDINGS" style to "Heading 2" (matching its sibling 4.x
     sections).
  3. Fix "4.6 Logical Database Design", which was mis-styled as Heading 1, to
     Heading 2 so it nests correctly under Chapter Four in the TOC.
  4. Ensure the document refreshes the TOC field on open.

Usage:
  python scripts/patch_chapter4_numbering_toc.py
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEFAULT_PATH = Path(__file__).resolve().parents[1] / "FINAL PROJECT_backup_revised.docx"


def style_exists(doc, name):
    return any(s.name == name for s in doc.styles)


def set_style(p, doc, style_name):
    if style_exists(doc, style_name):
        p.style = doc.styles[style_name]
        return True
    return False


def replace_prefix(p, old, new):
    """Replace the leading number `old` with `new` within the paragraph runs."""
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    # No run carried the text (e.g. text split oddly): rewrite first run.
    if p.runs and p.text.strip().startswith(old):
        p.runs[0].text = p.text.replace(old, new, 1)
        for r in p.runs[1:]:
            r.text = ""
        return True
    return False


def enable_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", type=Path, default=DEFAULT_PATH)
    args = parser.parse_args()
    if not args.path.exists():
        print(f"File not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(args.path))
    actions = []

    # 0) Appendix labels were numbered 5.x, colliding with Chapter Five (5.1/5.2/
    #    5.3). Relabel them as proper appendices so the numbering is logical and
    #    the TOC has no duplicate 5.x entries.
    appendix_renames = {
        "5. APPENDICES": "APPENDICES",
        "5.1 Appendix A": "Appendix A",
        "5.2 Appendix B": "Appendix B",
        "5.3 Appendix C: Data Collection Instruments": "Appendix C: Data Collection Instruments",
        "5.1.1 Time frame": "Appendix A.1 Time frame",
        "5.1.2 Gantt Chart": "Appendix A.2 Gantt Chart",
    }
    for p in doc.paragraphs:
        new = appendix_renames.get(p.text.strip())
        if new is not None and p.text.strip() != new:
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new)
            actions.append(f"relabelled appendix '{p.text.strip()[:30]}'")

    for p in doc.paragraphs:
        t = p.text.strip()

        # 1) Rename duplicate 4.8 Security Architecture -> 4.9
        if t == "4.8 Security Architecture":
            if replace_prefix(p, "4.8", "4.9"):
                actions.append("renamed 'Security Architecture' to 4.9")

        # 2) 4.8 Data Inputs -> Heading 2 (TOC visible)
        elif t == "4.8 Data Inputs" and p.style.name != "Heading 2":
            if set_style(p, doc, "Heading 2"):
                actions.append("'4.8 Data Inputs' set to Heading 2")

        # 3) 4.6 Logical Database Design -> Heading 2 (was Heading 1)
        elif t == "4.6 Logical Database Design" and p.style.name != "Heading 2":
            if set_style(p, doc, "Heading 2"):
                actions.append("'4.6 Logical Database Design' set to Heading 2")

        # 4) 4.5.1.2 is a 4th-level subsection mis-styled as Heading 2; align it
        #    with its siblings (4.5.1.1, 4.5.1.3) as Heading 4.
        elif t.startswith("4.5.1.2 ") and p.style.name != "Heading 4":
            if set_style(p, doc, "Heading 4"):
                actions.append(f"'{t[:40]}' set to Heading 4")

    # 5) Remove blank entries from the TOC: empty paragraphs that still carry a
    #    heading style produce blank lines in the table of contents.
    blanked = 0
    for p in doc.paragraphs:
        if not p.text.strip() and p.style.name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            if set_style(p, doc, "Normal"):
                blanked += 1
    if blanked:
        actions.append(f"cleared {blanked} empty heading-styled paragraph(s) from the TOC")

    enable_update_fields(doc)
    doc.save(str(args.path))

    if actions:
        print(f"Patched {args.path.name}:")
        for a in actions:
            print(f"  - {a}")
    else:
        print("No changes needed (already consistent).")
    print("  - TOC set to refresh on open")


if __name__ == "__main__":
    main()
