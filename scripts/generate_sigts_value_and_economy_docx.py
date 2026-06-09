"""Generate a Word doc capturing two SIGTS discussion responses:
1. What makes SIGTS unique/smart for Bwindi (realistic appraisal).
2. How SIGTS contributes to Uganda's economy without a fee/ads feature.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

OUT_PATH = Path(r"C:\Projects\SIGTS\DOCS\SIGTS_Value_and_Economic_Contribution.docx")


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, size=14 if level == 1 else 12, bold=True)


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), italic=italic)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        set_font(p.add_run(bold_prefix), bold=True)
    set_font(p.add_run(text))


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), italic=True)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, label in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = label
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run)
    doc.add_paragraph()


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("SIGTS: Realistic Value and Economic Contribution"), size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run("Smart Information Guide Tour System — Bwindi Impenetrable National Park, Uganda"), size=12)

    doc.add_paragraph()

    # =====================================================================
    # PART 1
    # =====================================================================
    heading(doc, "Part 1: What Makes SIGTS Unique and Smart (A Realistic Appraisal)", level=1)

    body(
        doc,
        "Bwindi Impenetrable National Park does not need SIGTS to operate. The Uganda Wildlife "
        "Authority (UWA), rangers, trackers, licensed guides, gorilla permits, safety briefings, and "
        "on-the-ground enforcement have sustained the park for decades. SIGTS is best understood as a "
        "digital operations and visitor-experience layer — valuable if adopted, but not essential in "
        "the way that permits, patrols, and veterinary care are essential."
    )

    heading(doc, "1.1 What is genuinely smart about SIGTS", level=2)

    bullet(doc, "Most parks juggle paper, radio, WhatsApp, PDFs, and separate permit systems. SIGTS unifies tourist information, guide workflows, sightings, geofence events, cultural content, feedback, and IT/admin tools in one platform. This integration design is the real contribution.", bold_prefix="One integrated platform: ")
    bullet(doc, "Poor forest connectivity, strict safety rules (7 m from gorillas, disease risk, behaviour near elephants), and UNESCO/community context are encoded in seeded species guides, safety copy, cultural narratives, and an offline-capable web/PWA approach.", bold_prefix="Built for Bwindi's real constraints: ")
    bullet(doc, "Park boundary, entry/exit events, and location logging are implemented with PostGIS, supporting accountability and future safety alerts.", bold_prefix="Geofence and location history: ")
    bullet(doc, "Quick sighting reports, verification, and a sightings feed can support biodiversity monitoring and guide coordination.", bold_prefix="Guide and sighting workflow: ")
    bullet(doc, "Publish/verify flows for community narratives fit Bwindi's conservation-plus-culture tourism story and differentiate it from apps that only show animal facts.", bold_prefix="Verified cultural storytelling: ")
    bullet(doc, "Tourist / guide / IT manager / admin roles with JWT authentication and protected admin APIs mirror how a real institution would deploy software.", bold_prefix="Role-based governance: ")
    bullet(doc, "The chat assistant grounds answers in the park catalogue (animals, locations, safety, culture) — a useful visitor aid when an external LLM is configured.", bold_prefix="AI tour help (when configured): ")

    heading(doc, "1.2 What SIGTS is NOT (avoiding overclaiming)", level=2)
    body(doc, "An honest report should distinguish implemented capability from aspiration:")
    add_table(
        doc,
        ["Claim often implied", "Reality"],
        [
            ["\"AI-powered park management\"", "AI is mostly tour-help Q&A; no trained wildlife models or AI ranger dispatch."],
            ["\"Offline-first mobile system\"", "Browser storage with partial sync; not a rugged field SQLite app."],
            ["\"Predictive analytics\"", "APIs exist; ML training pipelines and real congestion models are missing."],
            ["\"Emergency / ranger alerts\"", "Events are stored; no integrated dispatch to rangers."],
            ["\"Revenue protection\"", "No integration with UWA permit booking or payment systems."],
            ["\"Notifications\"", "Email/SMS need SMTP/Twilio; not configured in production."],
        ],
    )

    heading(doc, "1.3 What park management actually depends on today", level=2)
    for item in [
        "Gorilla permits and UWA rules (the economic and conservation backbone).",
        "Rangers and trackers locating habituated gorilla groups.",
        "Licensed guides enforcing safety and etiquette.",
        "Physical pre-trek briefings.",
        "Radio and phone coordination in the forest.",
        "Community lodges, transport, and local employment.",
    ]:
        bullet(doc, item)
    body(doc, "SIGTS does not replace any of these; at best it supports them digitally.")

    heading(doc, "1.4 Why Bwindi is still a strong sample for the project", level=2)
    bullet(doc, "Visitor safety, guide coordination, content delivery, and park oversight are genuine Bwindi needs.", bold_prefix="Real problem domain: ")
    bullet(doc, "Auth, geospatial data, content management patterns, sightings, analytics, admin, and offline patterns make a substantial systems project.", bold_prefix="Architectural breadth: ")
    bullet(doc, "A web/PWA plus cloud database avoids forcing every visitor to install an app, and B2G licensing to UWA/operators is a realistic revenue path.", bold_prefix="Uganda-relevant deployment: ")
    bullet(doc, "It demonstrates how a protected area could move from fragmented analog processes toward integrated digital operations.", bold_prefix="Modernization path: ")

    heading(doc, "1.5 Suggested positioning statement", level=2)
    quote(
        doc,
        "SIGTS is not a substitute for UWA's permit system, ranger patrols, or on-the-ground guide "
        "supervision. Its value lies in consolidating visitor education, guide workflows, sighting "
        "records, cultural content, and administrative oversight into one geospatially aware platform "
        "suited to Bwindi's connectivity and conservation context. Park management can operate without "
        "it today; adoption would improve information consistency, accountability, and visitor "
        "experience rather than enable core conservation functions."
    )

    body(
        doc,
        "In short, the uniqueness is the combination — geofence, roles, Bwindi-specific content, "
        "sightings, cultural verification, and a deployable web stack — not any single feature that no "
        "one else has ever built."
    )

    doc.add_page_break()

    # =====================================================================
    # PART 2
    # =====================================================================
    heading(doc, "Part 2: How SIGTS Contributes to Uganda's Economy Without a Fee or Advertising Feature", level=1)

    body(
        doc,
        "This is a fair challenge. SIGTS does not generate money directly — it has no payment gateway, "
        "no permit booking, no commission, and no advertising. Therefore any economic contribution is "
        "indirect: it works by improving the things that already generate money in Bwindi, rather than "
        "by collecting money itself."
    )

    heading(doc, "2.1 The money already exists — SIGTS only influences it", level=2)
    body(doc, "Bwindi's tourism economy runs on revenue streams that SIGTS does not touch:")
    for item in [
        "Gorilla permits (sold by UWA — high value, foreign currency).",
        "Park entry and activity fees.",
        "Guide and ranger services.",
        "Lodges, transport, food, crafts, and community enterprises.",
    ]:
        bullet(doc, item)
    body(
        doc,
        "Because SIGTS handles none of these transactions, it cannot \"make money\" in the cashier "
        "sense. What it can influence is how much of that existing money is earned, retained, and "
        "repeated. That is the only honest economic claim."
    )

    heading(doc, "2.2 The four realistic indirect channels", level=2)
    bullet(doc, "Clear safety guidance, strong species/cultural content, and a smoother guided experience raise satisfaction, recommendations, and repeat/referral demand — supporting the premium pricing UWA and operators already charge.", bold_prefix="A. Better visitor experience: ")
    bullet(doc, "Digital sighting logs, guide scheduling, tour accountability, and admin oversight reduce wasted effort and improve coordination, which is an economic gain even with no new revenue line.", bold_prefix="B. Operational efficiency: ")
    bullet(doc, "Geofence entry/exit events, location history, verified sightings, and audit logs create records that reduce leakage and informal activity and support better deployment decisions, protecting existing revenue.", bold_prefix="C. Data and accountability: ")
    bullet(doc, "Surfacing cultural narratives, community storytellers, lodges, and local context steers visitor attention and spending toward local enterprises, keeping more tourism money circulating locally.", bold_prefix="D. Local value-chain support: ")

    heading(doc, "2.3 Where the project itself earns money (the B2G model)", level=2)
    body(
        doc,
        "SIGTS earns through licensing the software to institutions, not through tourists. If a Ugandan "
        "software team builds and sells it, that is local ICT industry revenue, jobs, and skills:"
    )
    for item in [
        "Annual platform subscription per park/site (paid by UWA, an agency, or a concession operator).",
        "Implementation and onboarding fees.",
        "Training for guides and IT officers.",
        "Support contracts (standard / priority).",
        "Multi-park expansion to other Ugandan protected areas.",
    ]:
        bullet(doc, item)

    heading(doc, "2.4 The honest limitation", level=2)
    quote(
        doc,
        "SIGTS does not collect fees, process permits, or sell advertising, so it generates no direct "
        "tourism revenue. Its economic contribution is indirect: improving visitor experience and "
        "reputation (supporting demand and willingness to pay), increasing operational efficiency and "
        "accountability (protecting existing revenue and reducing cost), and strengthening local value "
        "chains (keeping tourism spending in the community). The platform itself earns revenue through "
        "a business-to-government licensing model rather than consumer payments."
    )

    heading(doc, "2.5 Optional future features for direct revenue", level=2)
    body(doc, "Not required for the current scope, but worth noting as future work:")
    for item in [
        "Permit / activity booking integration with UWA or a gateway (Flutterwave, MTN MoMo, Airtel Money) — booking fees or commission.",
        "A marketplace for lodges, transport, and crafts — referral/commission.",
        "Sponsored conservation content or controlled, ethical partner placements.",
        "A premium tourist tier (offline packs, premium guides, audio tours) — weighed carefully to avoid excluding low-budget or local visitors.",
    ]:
        bullet(doc, item)

    heading(doc, "2.6 Bottom line", level=2)
    body(
        doc,
        "Without a payment or advertising feature, SIGTS contributes to Uganda's economy indirectly — "
        "by making the existing tourism economy more attractive, efficient, accountable, and locally "
        "beneficial — and directly only through business-to-government licensing of the software to "
        "park institutions. It is an enabler and modernization tool, not a revenue-collection system."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
