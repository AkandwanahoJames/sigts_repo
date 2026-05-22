"""
Fix Chapter 4 section, table, and figure numbering in FINAL PROJECT_backup_revised - 2.docx.
Run: python scripts/fix_chapter4_numbering.py
"""
from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

DOC_PATH = Path(r"c:\Projects\SIGTS\FINAL PROJECT_backup_revised - 2.docx")

# Apply longest / most specific replacements first within each category.
HEADING_REPLACEMENTS = [
    ("4.1 Introduction", "4.0 Introduction"),
    ("4.2 System Study", "4.1 System Study"),
    ("4.2.1 Current Study", "4.1.1 Current Study"),
    ("4.2.2 Strength of the current system", "4.1.2 Strength of the current system"),
    ("4.2.3 Weakness of the current system", "4.1.3 Weakness of the current system"),
    ("4.3 Data Analysis and Findings", "4.2 Data Analysis and Findings"),
    ("4.3.1 Presentation of Findings", "4.2.1 Presentation of Findings"),
    ("4.3.2 Section A: Professional Background", "4.2.2 Section A: Professional Background"),
    ("4.3.2.1 Demographics of the Respondents", "4.2.2.1 Demographics of the Respondents"),
    ("4.3.2.2 Age Distribution of Respondents", "4.2.2.2 Age Distribution of Respondents"),
    ("4.3.2.3 Gender Distribution of Respondents", "4.2.2.3 Gender Distribution of Respondents"),
    (
        "4.3.2.3 Nationality  Distribution of Respondents",
        "4.2.2.4 Nationality Distribution of Respondents",
    ),
    ("4.3.3 Section B: Current Information Ecosystem", "4.2.3 Section B: Current Information Ecosystem"),
    (
        "4.3.2.3 Distribution of how currently Tourists obtain information while visiting the National Park ",
        "4.2.3.1 Distribution of how currently Tourists obtain information while visiting the National Park",
    ),
    (
        "4.3.2.3 Number of times Respondents often visit the National Park ",
        "4.2.3.2 Number of times Respondents often visit the National Park",
    ),
    (
        "4.3.2.3 Number of times Respondents Receive inconsistent, incomplete or contradictory information.",
        "4.2.3.3 Number of times Respondents receive inconsistent, incomplete or contradictory information",
    ),
    (
        "4.3.2.3 Current challenges experienced by tourists ",
        "4.2.3.4 Current challenges experienced by tourists",
    ),
    (
        "4.3.2.3 Number of   Respondents who would prefer a mobile system that provides tourism information and navigation servic",
        "4.2.3.5 Number of Respondents who would prefer a mobile system that provides tourism information and navigation services",
    ),
    (
        "4.3.4 Section C: Perceptions of a Web-Based Guide Tour System",
        "4.2.4 Section C: Perceptions of a Web-Based Guide Tour System",
    ),
    (
        "4.3.2.3 Devices Respondents would prefer to use during the tour ",
        "4.2.4.1 Devices Respondents would prefer to use during the tour",
    ),
    ("4.3.5\u00a0SECTION D: USER SATISFACTION AND IMPROVEMENT", "4.2.5 Section D: User Satisfaction and Improvement"),
    ("4.3.5 SECTION D: USER SATISFACTION AND IMPROVEMENT", "4.2.5 Section D: User Satisfaction and Improvement"),
    (
        "4.3.2.3 Respondents that recommend the Smart Information Guide Tour System t if implemented",
        "4.2.5.1 Respondents that recommend the Smart Information Guide Tour System if implemented",
    ),
    ("4.4 System Requirements", "4.3 System Requirements"),
    ("4.4.1 Software Requirements", "4.3.1 Software Requirements"),
    ("4.4.2 Hardware Requirements", "4.3.2 Hardware Requirements"),
    ("4.4.3 User Requirements", "4.3.3 User Requirements"),
    ("4.4.4 Functional Requirements", "4.3.4 Functional Requirements"),
    ("4.4.5 Non-Functional Requirements", "4.3.5 Non-Functional Requirements"),
    ("4.4.5.1 Usability", "4.3.5.1 Usability"),
    ("4.4.5.2 Performance", "4.3.5.2 Performance"),
    ("4.4.5.3 Security", "4.3.5.3 Security"),
    ("4.4.5.4 Scalability", "4.3.5.4 Scalability"),
    ("4.4.5.5 Reliability", "4.3.5.5 Reliability"),
    ("4.4.5.6 Compatibility", "4.3.5.6 Compatibility"),
    ("4.4.5.7 Maintainability", "4.3.5.7 Maintainability"),
    ("4.4.5.8 Data Integrity", "4.3.5.8 Data Integrity"),
    ("4.4.5.9 Offline Storage", "4.3.5.9 Offline Storage"),
    ("4.4.5.10 Battery Efficiency", "4.3.5.10 Battery Efficiency"),
    ("4.4.5.11 Concurrency", "4.3.5.11 Concurrency"),
    ("4.4.5.12 Localization", "4.3.5.12 Localization"),
    ("4.5 SYSTEM DESIGN", "4.4 System Design"),
    (
        "4.5.1 System Architecture for the Smart Information Guide Tour System (SIGTS)",
        "4.4.1 System Architecture for the Smart Information Guide Tour System (SIGTS)",
    ),
    ("4.5.1.1 Introduction to the Tiered Architecture", "4.4.1.1 Introduction to the Tiered Architecture"),
    ("4.5.1.2 Design for Smart Information Guide Tour System", "4.4.1.2 Design for Smart Information Guide Tour System"),
    (
        "4.5.1.3 Technical Architecture, Database, API, and Deployment Design",
        "4.4.1.3 Technical Architecture, Database, API, and Deployment Design",
    ),
    (
        "4.5.1.3.1 Artificial Intelligence and Natural Language Processing",
        "4.4.1.3.1 Artificial Intelligence and Natural Language Processing",
    ),
    ("4.5.1.3.2 Physical Database Schema", "4.4.1.3.2 Physical Database Schema"),
    ("4.5.1.3.3 REST API Endpoint Catalogue", "4.4.1.3.3 REST API Endpoint Catalogue"),
    ("4.5.1.4 Deployment Architecture", "4.4.1.4 Deployment Architecture"),
    ("4.5.2 Enhanced Entity-Relationship Diagram (EERD)", "4.4.2 Enhanced Entity-Relationship Diagram (EERD)"),
    ("4.5.3 Context Diagram (Level 0 DFD)", "4.4.3 Context Diagram (Level 0 DFD)"),
    ("4.5.4 Data Flow Diagram (DFD Level 1)", "4.4.4 Data Flow Diagram (DFD Level 1)"),
    ("4.5.5 Flowchart: Tourist App Workflow", "4.4.5 Flowchart: Tourist App Workflow"),
    ("4.5.6 Flowchart: Tour Guide Workflow", "4.4.6 Flowchart: Tour Guide Workflow"),
    ("4.5.7 Flowchart: IT Manager Workflow", "4.4.7 Flowchart: IT Manager Workflow"),
    ("4.6 Logical Database Design", "4.5 Logical Database Design"),
    ("4.7 Physical Database Design", "4.6 Physical Database Design"),
    ("4.8 Data Inputs", "4.7 Data Inputs"),
]

TABLE_FIGURE_REPLACEMENTS = [
    ("Table 4.9: Summary of software requirements", "Table 4.12: Summary of software requirements"),
    ("Table 4.10: Summary of hardware requirements", "Table 4.13: Summary of hardware requirements"),
    ("Table 4.16 presents the core physical tables", "Table 4.14 presents the core physical tables"),
    ("Table 4.16: Core physical database relations for SIGTS.", "Table 4.14: Core physical database relations for SIGTS."),
    ("Table 4.17 documents sixteen REST endpoints", "Table 4.15 documents sixteen REST endpoints"),
    ("Table 4.17: REST API endpoint catalogue for SIGTS.", "Table 4.15: REST API endpoint catalogue for SIGTS."),
    ("Figure 4.5: : Number of times", "Figure 4.5: Number of times"),
    ("Figure 4.6:Distribution of Respondents", "Figure 4.6: Distribution of Respondents"),
    (" Figure 4.2: Gender distribution of respondents.", "Figure 4.2: Gender distribution of respondents"),
    ("\xa0Figure 4.10: Devices Respondents would prefer", "Figure 4.10: Devices Respondents would prefer"),
    (
        "Figure 4.12: EERD diagram for SIGTS.",
        "Figure 4.13: EERD diagram for SIGTS (see also Section 4.4.2).",
    ),
    ("Figure 4.11: System architecture for SIGTS.", "Figure 4.12: System architecture for SIGTS."),
    (
        "Figure 4.12: EERD of the Smart Information Guide Tour System of the Smart Information Guide Tour System",
        "Figure 4.13: EERD of the Smart Information Guide Tour System",
    ),
    (
        "Figure 4.13: Context diagram (Level 0 DFD) of the Smart Information Guide Tour System\u00a0",
        "Figure 4.14: Context diagram (Level 0 DFD) of the Smart Information Guide Tour System",
    ),
    ("Figure 4.14: DFD Level 1\u00a0", "Figure 4.15: DFD Level 1"),
    ("Figure 4.15: Tourist app workflow workflow\u00a0", "Figure 4.16: Tourist app workflow"),
    ("Figure 4.16: Tour guide workflow workflow\u00a0", "Figure 4.17: Tour guide workflow"),
    ("Figure 4.17: IT manager workflow workflow\u00a0", "Figure 4.18: IT manager workflow"),
    ("Table 4.12: Physical design for tourists", "Table 4.16: Physical design for tourists"),
    ("Table 4.14: Physical design for wildlife sightings", "Table 4.19: Physical design for wildlife sightings"),
    ("Table 4.15: Physical design for notifications", "Table 4.20: Physical design for notifications"),
    ("Figure 4.26: IT manager dashboard for SIGTS", "Figure 4.27: IT manager dashboard for SIGTS"),
    ("Figure 4.25: Cultural narratives page for SIGTS", "Figure 4.26: Cultural narratives page for SIGTS"),
    ("Figure 4.24: Tour guide dashboard for SIGTS", "Figure 4.25: Tour guide dashboard for SIGTS"),
    ("Figure 4.23: Wildlife sightings page for SIGTS", "Figure 4.24: Wildlife sightings page for SIGTS"),
    ("Figure 4.22: Interactive map view for SIGTS", "Figure 4.23: Interactive map view for SIGTS"),
    ("Figure 4.21: Wildlife information page for SIGTS", "Figure 4.22: Wildlife information page for SIGTS"),
    ("Figure 4.20: Tourist dashboard (home screen) for SIGTS", "Figure 4.21: Tourist dashboard (home screen) for SIGTS"),
    ("Figure 4.19: Registration form for SIGTS", "Figure 4.20: Registration form for SIGTS"),
    ("Figure 4.18: Login form for SIGTS", "Figure 4.19: Login form for SIGTS"),
]

BODY_CROSS_REFS = [
    ("Chapter Four", "Chapter Four"),
    ("section 4.4", "section 4.3"),
    ("Section 4.5", "Section 4.4"),
    ("4.5.1", "4.4.1"),
]


def para_text(p) -> str:
    return "".join(r.text for r in p.runs)


def set_para_text(p, text: str):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def apply_replacements(text: str, pairs: list[tuple[str, str]]) -> str:
    out = text
    for old, new in pairs:
        if old in out:
            out = out.replace(old, new)
    return out


def find_chapter_bounds(doc: Document) -> tuple[int, int] | None:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = para_text(p).strip()
        if re.search(r"^CHAPTER FOUR", t, re.I):
            start = i
        if start is not None and re.search(r"^CHAPTER FIVE", t, re.I):
            end = i
            break
    if start is None:
        return None
    return start, end or len(doc.paragraphs)


def fix_heading_styles(doc: Document, start: int, end: int):
    """Normalize heading levels for design subsections and database sections."""
    targets = {
        "4.4.1.2 Design for Smart Information Guide Tour System": 4,
        "4.5 Logical Database Design": 2,
        "4.6 Physical Database Design": 2,
        "4.7 Data Inputs": 2,
    }
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = para_text(p).strip()
        if t not in targets:
            continue
        level = targets[t]
        try:
            p.style = doc.styles[f"Heading {level}"]
        except KeyError:
            pass


def fix_first_guide_table(doc: Document, start: int, end: int):
    """Rename duplicate 4.13a captions to Tables 4.17 and 4.18."""
    seen_13a = False
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = para_text(p)
        if re.search(r"Table 4\.13a:\s*Physical design for tour guides", t, re.I):
            if not seen_13a:
                set_para_text(p, "Table 4.17: Physical design for tour guides")
                seen_13a = True
            else:
                set_para_text(p, "Table 4.18: Physical design for tour guides (continued)")


def fix_duplicate_eerd_caption(doc: Document, start: int, end: int):
    """Remove duplicate Figure 4.13 under architecture; keep canonical caption in 4.4.2."""
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = para_text(p)
        if "Figure 4.13: EERD diagram for SIGTS (see also Section 4.4.2)" in t:
            set_para_text(
                p,
                "The enhanced entity-relationship diagram is presented in Section 4.4.2 (Figure 4.13).",
            )
            break


def main():
    if not DOC_PATH.is_file():
        print(f"Missing {DOC_PATH}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(DOC_PATH))
    bounds = find_chapter_bounds(doc)
    if not bounds:
        print("Chapter Four not found", file=sys.stderr)
        sys.exit(1)

    start, end = bounds
    all_pairs = HEADING_REPLACEMENTS + TABLE_FIGURE_REPLACEMENTS
    changed = 0

    for i in range(start, end):
        p = doc.paragraphs[i]
        old = para_text(p)
        if not old.strip():
            continue
        new = apply_replacements(old, all_pairs)
        new = apply_replacements(new, BODY_CROSS_REFS)
        if new != old:
            set_para_text(p, new)
            changed += 1

    fix_first_guide_table(doc, start, end)
    fix_duplicate_eerd_caption(doc, start, end)
    fix_heading_styles(doc, start, end)

    doc.save(str(DOC_PATH))
    print(f"Updated {changed} paragraphs in Chapter 4 ({DOC_PATH.name})")

    # Verification summary
    print("\n--- Chapter 4 outline after fix ---")
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = para_text(p).strip()
        style = p.style.name if p.style else ""
        if re.match(r"^4\.\d", t) or re.match(r"^(Table|Figure) 4\.\d", t, re.I):
            print(f"  {t[:90]}")


if __name__ == "__main__":
    main()
