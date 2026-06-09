"""Generate SIGTS Test Plan and Environment Test Matrix as Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

OUT_PATH = Path(r"C:\Projects\SIGTS\DOCS\SIGTS_Test_Plan_and_Environment_Matrix.docx")


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, size=14 if level == 1 else 12, bold=True)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_font(run)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_font(run)


def add_matrix_table(doc):
    """Environment test matrix as a Word table."""
    heading(doc, "Environment Test Matrix (Summary Table)", level=2)
    body(
        doc,
        "The table below summarizes what to test in each environment, which tools to use, "
        "and the release gate for that stage."
    )

    headers = [
        "Environment",
        "Purpose",
        "Data",
        "Tests required",
        "Tools",
        "Release gate",
    ]
    rows = [
        [
            "Local development",
            "Fast feedback during implementation",
            "Local PostgreSQL + seed/demo data",
            "Unit, integration, API regression subset",
            "Jest, Postman/Newman, browser devtools",
            "All unit/integration tests pass before merge",
        ],
        [
            "Hosted pre-release (staging-like)",
            "Validate deployment and integration before production",
            "Hosted DB with non-sensitive test data",
            "Full API regression, UAT, notification checks, security baseline",
            "Newman, smoke script, SCA/secret scan, OWASP ZAP (lite)",
            "All critical scenarios pass; no blocker defects",
        ],
        [
            "Production",
            "Confirm live health and critical paths after deploy",
            "Live production data",
            "Smoke tests only (non-destructive)",
            "PowerShell/curl, Invoke-RestMethod, Vercel logs",
            "Health healthy; login + protected endpoint OK; no high-severity log errors",
        ],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, label in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = label
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True)

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run)

    doc.add_paragraph()


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SIGTS Test Plan and Environment Test Matrix")
    set_font(r, size=14, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Smart Information Guide Tour System — Bwindi Impenetrable National Park")
    set_font(r2, size=12)

    doc.add_paragraph()

    heading(doc, "1. Purpose", level=1)
    body(
        doc,
        "This document defines a practical test plan for SIGTS and a clear environment matrix "
        "to verify quality before and after releases. It is designed to ensure account registration "
        "and login remain reliable, protected workflows are secure and role-correct, production "
        "deployments are validated quickly, and regressions are detected early."
    )

    heading(doc, "2. Scope", level=1)
    body(doc, "Included:")
    for item in [
        "Frontend web/PWA behavior for core user journeys.",
        "Backend API behavior, database interactions, auth, and role access.",
        "Notification behavior (email/SMS configured vs not configured).",
        "Production smoke checks and post-deploy validation.",
    ]:
        bullet(doc, item)
    body(doc, "Excluded:")
    for item in [
        "Deep penetration testing (dedicated security assessments).",
        "Full load/performance certification at enterprise scale (separate exercise).",
    ]:
        bullet(doc, item)

    heading(doc, "3. Test Objectives", level=1)
    for item in [
        "Validate critical user journeys for Tourist, Guide, IT Manager, and Admin.",
        "Verify new account creation can always log in.",
        "Confirm health, DB connectivity, and protected endpoint access in production.",
        "Ensure failures degrade gracefully (for example missing SMTP/Twilio).",
    ]:
        bullet(doc, item)

    heading(doc, "4. Test Types and Minimum Coverage", level=1)

    heading(doc, "4.1 Unit tests", level=2)
    bullet(doc, "Auth helpers, normalization, notification reason mapping, guard logic.")
    bullet(doc, "Target: at least 80% coverage on critical auth/notification modules.")

    heading(doc, "4.2 Integration tests", level=2)
    bullet(doc, "Register → login → token use → protected endpoint.")
    bullet(doc, "Duplicate account conflict handling.")
    bullet(doc, "Role-based access checks for admin endpoints.")

    heading(doc, "4.3 API regression tests", level=2)
    bullet(doc, "Core routes: auth, users/profile, animals/locations, admin/system-health.")
    bullet(doc, "Negative cases: wrong password, missing token, role mismatch.")

    heading(doc, "4.4 User acceptance testing (UAT)", level=2)
    bullet(doc, "Tourist, Guide, IT Manager scripted tasks.")
    bullet(doc, "Focus on usability and workflow completeness.")

    heading(doc, "4.5 Smoke tests (post-deploy)", level=2)
    bullet(doc, "/api/health, demo login, protected endpoint, new registration/login.")
    bullet(doc, "Pass/fail within 5–10 minutes per deployment.")

    heading(doc, "4.6 Security baseline tests", level=2)
    bullet(doc, "JWT and authorization checks, input validation, dependency scan, secret scan.")
    bullet(doc, "Suggested tools: OWASP ZAP, Semgrep, npm audit, Gitleaks, Postman abuse-case suite.")

    heading(doc, "5. Entry and Exit Criteria", level=1)
    body(doc, "Entry criteria:")
    for item in [
        "Environment is reachable.",
        "Required migrations applied.",
        "Seed/demo data available (for non-production smoke/UAT where needed).",
        "Test credentials and base URLs are available.",
    ]:
        bullet(doc, item)
    body(doc, "Exit criteria:")
    for item in [
        "No blocker or critical defects open.",
        "All smoke tests pass.",
        "Core auth/account scenarios pass in target environment.",
        "Known medium/low issues documented with owner and fix timeline.",
    ]:
        bullet(doc, item)

    heading(doc, "6. Environment Test Matrix (Detailed)", level=1)

    envs = [
        (
            "Local Development",
            [
                ("Purpose", "Fast developer feedback during implementation."),
                ("Data", "Local database and local seed/demo data."),
                ("Tests required", "Unit tests; integration tests; API regression subset."),
                ("Tools", "Node test runner/Jest; Postman/Newman; browser developer tools."),
                ("Gate", "No broken unit/integration tests before merge."),
            ],
        ),
        (
            "Hosted Pre-Release (Staging-like)",
            [
                ("Purpose", "Validate deployment behavior and realistic integration before production."),
                ("Data", "Hosted database with non-sensitive test data."),
                (
                    "Tests required",
                    "Full API regression; role-based UAT; notification behavior checks; security baseline scan.",
                ),
                (
                    "Tools",
                    "Newman; scripted smoke runner; security scanners (SCA, secret scan, DAST-lite).",
                ),
                ("Gate", "All critical scenarios pass; no blocker defects."),
            ],
        ),
        (
            "Production",
            [
                ("Purpose", "Confirm service health and critical functionality after deployment."),
                ("Data", "Live production data."),
                (
                    "Tests required",
                    "Smoke tests only (safe, non-destructive); health and auth checks; log review.",
                ),
                ("Tools", "PowerShell smoke script; curl/Invoke-RestMethod; Vercel deployment logs."),
                (
                    "Gate",
                    "/api/health healthy with database connected; login and protected endpoint succeed; "
                    "no high-severity runtime errors in deployment logs.",
                ),
            ],
        ),
    ]
    for env_name, fields in envs:
        heading(doc, f"Environment: {env_name}", level=2)
        for label, text in fields:
            p = doc.add_paragraph()
            r0 = p.add_run(f"{label}: ")
            set_font(r0, bold=True)
            r1 = p.add_run(text)
            set_font(r1)

    add_matrix_table(doc)

    heading(doc, "7. Core Scenario Matrix (Must Always Pass)", level=1)

    scenarios = [
        (
            "Auth and accounts",
            [
                "Register new user → login succeeds.",
                "Existing user login succeeds with correct credentials.",
                "Incorrect password returns controlled 401 response.",
                "Inactive/deactivated account cannot log in.",
            ],
        ),
        (
            "Role and access",
            [
                "Tourist cannot access admin-only operations.",
                "IT/Admin can access system-health and management endpoints.",
                "Protected endpoints reject missing or invalid tokens.",
            ],
        ),
        (
            "System health",
            [
                "Health endpoint reports status healthy in production.",
                "Database state reports connected.",
            ],
        ),
        (
            "Notification behavior",
            [
                "When SMTP/Twilio not configured: registration succeeds and reason codes are returned.",
                "When SMTP/Twilio configured: delivery attempts succeed or errors are explicit.",
            ],
        ),
    ]
    for group, items in scenarios:
        body(doc, group + ":")
        for item in items:
            bullet(doc, item)

    heading(doc, "8. Defect Severity and Response", level=1)
    severities = [
        ("Blocker", "Login broken, registration broken, health/DB down.", "Immediate hotfix; stop release."),
        ("High", "Role bypass, token auth failure in core flows.", "Fix before release."),
        ("Medium", "Non-core endpoint issues, partial UX regressions.", "Schedule in next sprint."),
        ("Low", "Cosmetic or minor messaging issues.", "Backlog."),
    ]
    for sev, examples, action in severities:
        p = doc.add_paragraph()
        set_font(p.add_run(f"{sev}: "), bold=True)
        set_font(p.add_run(f"{examples} Action: {action}"))

    heading(doc, "9. Execution Ownership", level=1)
    owners = [
        ("Developers", "Unit/integration/API regression in local and pre-release."),
        ("QA / Test Lead", "Regression packs, UAT coordination, defect triage."),
        ("DevOps / Release Owner", "Deployment verification, production smoke, log inspection."),
        ("Product Owner / Stakeholders", "UAT sign-off and release approval."),
    ]
    for role, duty in owners:
        p = doc.add_paragraph()
        set_font(p.add_run(f"{role}: "), bold=True)
        set_font(p.add_run(duty))

    heading(doc, "10. Release-Day Smoke Checklist", level=1)
    for item in [
        "Confirm deployment completed successfully.",
        "Run production health check (GET /api/health).",
        "Execute demo account login and one protected endpoint call.",
        "Execute one fresh registration and login validation (when safe).",
        "Verify logs show no high-severity runtime failures.",
        "Record results in release note with timestamp and owner.",
    ]:
        bullet(doc, item)

    heading(doc, "11. Suggested Artifacts to Maintain", level=1)
    for item in [
        "Smoke test script output (per deployment).",
        "API regression report (Newman or equivalent).",
        "UAT checklist and sign-off record.",
        "Defect tracker snapshot for release decision.",
        "Security scan summary (dependencies and secrets).",
    ]:
        bullet(doc, item)

    body(doc, "Prepared for SIGTS release governance and continuous quality assurance.")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
