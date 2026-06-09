from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT_PATH = Path(r"C:\Projects\SIGTS\SIGTS_Technology_and_Business_Model_Uganda.docx")


def set_run_style(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_style(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_run_style(run, size=12)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_style(run, size=12)


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SIGTS: Technology Stack, Business Model, and Revenue Impact for Uganda")
    set_run_style(r, size=14, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Prepared for strategic and operational use")
    set_run_style(r2, size=12)

    add_heading(doc, "1. Overview", level=1)
    add_paragraph(
        doc,
        "The Smart Information Guide Tour System (SIGTS) is a digital platform for park tourism "
        "operations, visitor guidance, safety, and management oversight. It is designed for protected "
        "areas such as Bwindi and supports tourists, guides, and IT/administrative teams through a "
        "single web platform."
    )

    add_heading(doc, "2. Technologies Used to Build SIGTS", level=1)
    add_heading(doc, "2.1 Frontend Layer", level=2)
    add_bullet(doc, "Web application delivered from Vercel as a static Progressive Web App (PWA).")
    add_bullet(doc, "JavaScript-based user interfaces for tourists, guides, and IT/admin users.")
    add_bullet(doc, "Runtime configuration pattern for production API routing and environment setup.")
    add_bullet(doc, "Map and geospatial interaction components for routes, boundaries, and locations.")

    add_heading(doc, "2.2 Backend/API Layer", level=2)
    add_bullet(doc, "Node.js + Express API architecture.")
    add_bullet(doc, "JWT-based authentication and role-based access controls.")
    add_bullet(doc, "Operational modules for tours, sightings, geofence events, analytics, and intranet use.")
    add_bullet(doc, "Production deployment as serverless API endpoints on Vercel.")

    add_heading(doc, "2.3 Data and Infrastructure Layer", level=2)
    add_bullet(doc, "PostgreSQL database hosted on Supabase.")
    add_bullet(doc, "Schema migrations and seeded operational/demo data scripts.")
    add_bullet(doc, "Geospatial capability through PostGIS-backed design.")
    add_bullet(doc, "Production environment configuration with secure variables and health monitoring.")

    add_heading(doc, "2.4 DevOps and Operations", level=2)
    add_bullet(doc, "Version-controlled monorepo with frontend/backend workspaces.")
    add_bullet(doc, "Environment-based deployment workflow for production.")
    add_bullet(doc, "Health endpoint and role-based smoke testing for post-deployment validation.")
    add_bullet(doc, "Operational scripts for account reset, migration prep, and service verification.")

    add_heading(doc, "3. Business Model for SIGTS", level=1)
    add_paragraph(
        doc,
        "SIGTS is best positioned as a Business-to-Government (B2G) and Business-to-Institution (B2I) "
        "platform for protected-area management rather than a direct consumer-paid app."
    )

    add_heading(doc, "3.1 Core Model", level=2)
    add_bullet(doc, "Primary customer: Uganda Wildlife Authority, park authorities, tourism agencies, or concession operators.")
    add_bullet(doc, "Revenue basis: annual platform subscription per park/site.")
    add_bullet(doc, "Commercial scope: software license + support + updates + operational analytics.")

    add_heading(doc, "3.2 Additional Revenue Streams", level=2)
    add_bullet(doc, "Implementation fees: onboarding, customization, and institutional rollout.")
    add_bullet(doc, "Training and capacity building: staff training for guides and IT officers.")
    add_bullet(doc, "Support tiers: standard, priority, and mission-critical support contracts.")
    add_bullet(doc, "Multi-site expansion: deployment to additional parks or conservation areas.")

    add_heading(doc, "4. How SIGTS Makes Money for Uganda", level=1)
    add_paragraph(
        doc,
        "SIGTS supports national income generation by increasing tourism efficiency, improving visitor "
        "experience, and strengthening revenue integrity in park operations."
    )

    add_heading(doc, "4.1 Direct Economic Contribution", level=2)
    add_bullet(doc, "Improves park service quality, supporting higher visitor willingness to pay for permits and guided services.")
    add_bullet(doc, "Reduces revenue leakage through better digital records, route tracking, and operational accountability.")
    add_bullet(doc, "Enables tiered digital services for institutions and operators under formal contracts.")

    add_heading(doc, "4.2 Indirect Economic Contribution", level=2)
    add_bullet(doc, "Increases tourist satisfaction, promoting repeat visits and positive international reputation.")
    add_bullet(doc, "Supports growth of local value chains: guides, transport, lodging, food, crafts, and community enterprises.")
    add_bullet(doc, "Strengthens conservation sustainability by linking improved operations to stronger long-term tourism revenue.")

    add_heading(doc, "4.3 National-Level Financial Effects", level=2)
    add_bullet(doc, "Higher foreign exchange inflows from inbound tourism.")
    add_bullet(doc, "Increased tax collection from tourism-linked businesses.")
    add_bullet(doc, "Better policy planning through aggregated operational analytics and performance data.")

    add_heading(doc, "5. Recommended Commercialization Path", level=1)
    add_bullet(doc, "Phase 1: Stabilize production operations and publish measurable service KPIs.")
    add_bullet(doc, "Phase 2: Formalize annual licensing and support contracts for operating institutions.")
    add_bullet(doc, "Phase 3: Scale to multiple protected areas with standardized deployment packages.")
    add_bullet(doc, "Phase 4: Introduce premium analytics and cross-park benchmarking as paid modules.")

    add_heading(doc, "6. Conclusion", level=1)
    add_paragraph(
        doc,
        "SIGTS combines practical tourism operations technology with a scalable institutional business model. "
        "If implemented across parks with clear governance and service contracts, it can generate direct platform "
        "revenue, increase tourism-sector income, and strengthen Uganda's broader conservation economy."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
