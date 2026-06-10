"""
Expand section 6.4 Recommendations with UWA, government, and business-model guidance.

Usage:
  python scripts/patch_recommendations_business_model.py
"""
from __future__ import annotations

import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DEFAULT_PATH = Path(__file__).resolve().parents[1] / "FINAL PROJECT_backup_revised.docx"
SUBHEAD_STYLE = "MY HEANDINGS"


def replace_paragraph_text(p, new_text: str):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def style_exists(doc, name: str) -> bool:
    return any(s.name == name for s in doc.styles)


def insert_paragraph_before(anchor: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    new_para = Paragraph(new_p, anchor._parent)
    if text:
        replace_paragraph_text(new_para, text)
    if style_name and style_exists(anchor.part.document, style_name):
        new_para.style = anchor.part.document.styles[style_name]
    return new_para


def insert_block_before(anchor: Paragraph, blocks: list[tuple[str, str | None]]):
    cursor = anchor
    for text, style in reversed(blocks):
        cursor = insert_paragraph_before(cursor, text, style)
    return cursor


INTRO = (
    "The recommendations below move SIGTS from academic prototype toward sustained park service. "
    "They are grouped for the Uganda Wildlife Authority (UWA) and Bwindi park management, for "
    "government and partner agencies, and for a financing model that links verified information "
    "delivery to national tourism revenue goals under NDPIV."
)

UWA_HEAD = "6.4.1 Recommendations for the Uganda Wildlife Authority and Park Management"
UWA_BODY = (
    "UWA should establish a named SIGTS content governance unit at Bwindi with stewards responsible "
    "for species records, cultural narratives, permit guidance, and safety messaging. All material "
    "published through the platform should carry explicit UWA approval status, consistent with "
    "existing authority over gorilla trekking rules and community revenue-sharing communication. "
    "A sector-by-sector pilot (Buhoma, Ruhija, Rushaga, Nkuringo) with licensed guides and lodge "
    "partners should precede park-wide rollout, using the Appendix E validation checklist as the "
    "acceptance gate. Over time, visitor-flow analytics, feedback, and sighting records collected "
    "in SIGTS should feed UWA operational reporting rather than remaining isolated from ranger "
    "briefings and sector management. Guide professionalization through USAGA-aligned training "
    "should position SIGTS as a supplement to live interpretation during gorilla encounters, not a "
    "substitute. Finally, UWA should maintain a quarterly content refresh calendar tied to seasonal "
    "trekking conditions, official tariff updates, and community partnership agreements, because "
    "outdated or inconsistent information was the primary complaint identified in Chapter 4."
)

GOV_HEAD = "6.4.2 Recommendations for Government and Inter-Agency Partners"
GOV_BODY = (
    "The Ministry of Tourism, Wildlife and Antiquities (MoTWA), Uganda Tourism Board (UTB), and "
    "National Information Technology Authority–Uganda (NITA-U) should treat SIGTS as a strategic "
    "enabler of NDPIV tourism–ICT convergence rather than a park-only information technology "
    "project. MoTWA should ring-fence digital tourism transformation funding for intranet hosting, "
    "sector connectivity, and cross-park replication once the Bwindi pilot succeeds—comparable to "
    "regional investments in integrated destination platforms. UTB should link UWA-verified SIGTS "
    "content with the Explore Uganda ecosystem so pre-visit promotion and on-site park facts share "
    "one authoritative narrative, directly addressing the dissatisfaction linked to limited site "
    "information in national surveys. NITA-U should host or certify SIGTS infrastructure against "
    "national e-government security standards and the Data Protection and Privacy Act (2019), "
    "ensuring feedback and usability data are processed lawfully. Government should also broker "
    "public–private connectivity partnerships with licensed telecommunications operators to extend "
    "reliable backhaul to sector offices, converting uneven field connectivity from a technical risk "
    "into shared infrastructure investment aligned with the UWA–UTB–NITA-U digital partnership."
)

BIZ_HEAD = "6.4.3 Sustainable Business Model and Institutional Financing"
BIZ_BODY_1 = (
    "SIGTS should not depend indefinitely on ad hoc project grants. The recommended model is a "
    "public-anchored park service whose business case rests on information quality, visitor "
    "satisfaction, and protected permit revenue rather than direct tourist charges for basic facts. "
    "International tourism receipts reached approximately USD 1.28 billion nationally in 2024; "
    "within that envelope, UWA should fund core content stewardship and intranet operations from "
    "conservation and permit income, treating SIGTS like wayfinding signage or ranger briefing "
    "centres—destination infrastructure that protects brand integrity and reduces misinformation "
    "liability."
)
BIZ_BODY_2 = (
    "Marginal scale costs can be recovered through structured partnerships: licensed tour operators "
    "and lodges contribute modest annual enablement fees in exchange for authenticated guide "
    "accounts and pre-visit content packs; cross-park licensing spreads fixed information "
    "technology costs when modules deploy to Queen Elizabeth, Murchison Falls, and Kibale under a "
    "single UWA digital services framework; and operational analytics (congestion prediction, "
    "staffing recommendations) justify internal cost recovery by reducing waste without paywalls "
    "on public park information. Repeat visitation, Net Promoter Score, and System Usability Scale "
    "results from user acceptance testing should be tracked as return-on-investment indicators for "
    "MoTWA and UWA. Cultural narrative modules should transparently present official community fee "
    "and revenue-sharing arrangements, strengthening trust without commercialising sensitive "
    "content. Direct pay-per-download monetisation of core facts is discouraged because gorilla "
    "trekking visitors already pay substantial permit fees; adding friction at the information "
    "layer would undermine the equitable public-good purpose SIGTS was designed to serve."
)


def patch(doc: Document) -> list[str]:
    actions: list[str] = []
    rec_heading = None
    rec_body = None
    concl_heading = None

    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "6.4 Recommendations":
            rec_heading = p
        elif t == "6.5 Conclusion":
            concl_heading = p
        elif rec_heading and rec_body is None and t.startswith("Park management and the development team"):
            rec_body = p

    if not rec_heading or not rec_body or not concl_heading:
        return ["ERROR: could not locate section 6.4 paragraphs"]

    if "6.4.1 Recommendations for the Uganda Wildlife Authority" in concl_heading.text:
        return ["section 6.4 already expanded"]

    # Detect prior patch
    for p in doc.paragraphs:
        if p.text.strip().startswith("6.4.1 Recommendations for the Uganda Wildlife Authority"):
            return ["section 6.4 already expanded"]

    replace_paragraph_text(rec_body, INTRO)
    actions.append("replaced 6.4 intro")

    sub_style = SUBHEAD_STYLE if style_exists(doc, SUBHEAD_STYLE) else None
    blocks = [
        (UWA_HEAD, sub_style),
        (UWA_BODY, None),
        (GOV_HEAD, sub_style),
        (GOV_BODY, None),
        (BIZ_HEAD, sub_style),
        (BIZ_BODY_1, None),
        (BIZ_BODY_2, None),
    ]
    insert_block_before(concl_heading, blocks)
    actions.append("inserted 6.4.1–6.4.3 recommendations")
    return actions


def main():
    path = DEFAULT_PATH
    if not path.exists():
        print(f"Missing: {path}", file=sys.stderr)
        sys.exit(1)

    backup = path.with_suffix(f".pre-recs-{date.today().isoformat()}.docx")
    shutil.copy2(path, backup)

    doc = Document(str(path))
    actions = patch(doc)
    if actions and actions[0].startswith("ERROR"):
        print(actions[0], file=sys.stderr)
        sys.exit(1)

    doc.save(str(path))
    print(f"Backup: {backup.name}")
    for a in actions:
        print(f"  - {a}")
    print("Update the Table of Contents in Word (References -> Update Table).")


if __name__ == "__main__":
    main()
