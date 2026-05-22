"""Generate SIGTS report Chapters 5 and 6 as a standalone Word document."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_PATH = r"c:\Projects\SIGTS\SIGTS_Chapters_5_and_6_Report.docx"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            set_font(run, 12 + (2 if level == 1 else 0), bold=True)

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_font(run)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SMART INFORMATION GUIDE TOUR SYSTEM (SIGTS)\n")
    set_font(r, 14, bold=True)
    r2 = t.add_run("Bwindi Impenetrable National Park\n\n")
    set_font(r2)
    r3 = t.add_run("Chapters Five and Six — Project Report Extract")
    set_font(r3)

    doc.add_page_break()

    heading("CHAPTER FIVE: SYSTEM IMPLEMENTATION, TESTING AND VALIDATION", 1)
    body(
        "5.0 Introduction\n"
        "This chapter documents how the Smart Information Guide Tour System (SIGTS) was "
        "translated from design into a working application. It describes the technologies "
        "employed, the principal functions delivered to tourists, field guides, and park "
        "administrators, and the procedures used to test and validate data handled by the system."
    )

    heading("5.1 System Implementation", 2)
    body(
        "Implementation followed a three-tier arrangement: a browser-based client served from "
        "the project frontend, a Node.js application programming interface backed by PostgreSQL "
        "with PostGIS extensions, and supporting services for email, optional large-language-model "
        "tour assistance, and real-time notifications. Development was carried out in modular "
        "stages so that authentication, catalogue content, mapping, sightings, guide scheduling, "
        "and administrative analytics could be integrated incrementally."
    )

    heading("5.1.1 System technologies", 3)
    body(
        "The client layer uses standards-based web technologies. HyperText Markup Language (HTML) "
        "defines page structure; Cascading Style Sheets control presentation; and JavaScript "
        "implements navigation, form handling, and communication with the server. Leaflet was "
        "adopted for interactive mapping of gates, trails, and points of interest within the park "
        "boundary. A service worker caches selected assets so that tourists retain access to "
        "catalogue material when connectivity is weak inside the forest."
    )
    body(
        "The server layer is built on Express.js running on Node.js. Representational State "
        "Transfer (REST) endpoints expose park content, user accounts, sightings, tour sessions, "
        "cultural narratives, geofence checks, feedback, and analytics. JSON Web Tokens secure "
        "authenticated sessions, with refresh-token rotation and configurable idle timeout. "
        "PostgreSQL stores relational data; PostGIS supports spatial queries for boundaries, "
        "visitor location context, and proximity to named places."
    )
    body(
        "Technologies were selected for availability, maintainability within the student project "
        "timeline, and alignment with park operations. Open-source components reduced licensing "
        "cost while allowing the team to deploy the stack on modest hardware suitable for "
        "demonstration and pilot use at Bwindi Impenetrable National Park."
    )

    heading("5.2 System Functionality", 2)
    heading("5.2.1 Authentication and role-based access", 3)
    body(
        "The system distinguishes tourists, registered guides, and information-technology managers. "
        "Visitors authenticate before accessing personalised dashboards, saved bookmarks, and sighting "
        "submission. Guides receive schedules, shift controls, and guest lists tied to assigned treks. "
        "IT managers access predictive analytics, backup controls, cultural narrative publishing, "
        "and operational summaries. Registration, password recovery, and email verification flows "
        "were implemented so that accounts could be provisioned without manual database edits."
    )
    heading("5.2.2 Map guidance and location services", 3)
    body(
        "The map module combines park boundary geometry, location records, trek routes, and recent "
        "wildlife sightings. Users may search for gates and sectors, measure approximate distances, "
        "and receive contextual notes when geolocation is available. Geofence logic supports "
        "demonstrations of on-site versus off-site access policies required for certain features "
        "during the pilot."
    )
    heading("5.2.3 Visitor information, culture, and tour assistance", 3)
    body(
        "Tourists browse a species catalogue enriched with conservation status, habitat notes, and "
        "safety guidance drawn from park-approved content. Cultural narratives present community "
        "stories with publishing controls managed by authorised staff. An integrated tour-help "
        "facility answers questions using curated rules when offline and, when configured, a "
        "grounded language model that retrieves answers from live database records rather than "
        "unverified external sources."
    )
    heading("5.2.4 Guide operations and administrative analytics", 3)
    body(
        "Guides record sightings, review preparation checklists, and monitor alerts for unusual "
        "observations. The IT dashboard aggregates visitor-flow indicators, satisfaction scores, "
        "sightings trends, congestion forecasts, and operational status for scheduled reports, "
        "model retraining jobs, and backup activity. These facilities support day-to-day park "
        "management decisions described in earlier chapters of this report."
    )

    heading("5.3 System Testing", 2)
    body(
        "Testing was planned to confirm that each stakeholder role could complete essential tasks "
        "without critical defects. Both manual inspection and structured walkthroughs were applied "
        "before the demonstration build was released for academic review."
    )
    heading("5.3.1 User testing", 3)
    body(
        "Representative users were invited to perform scripted tasks while observers recorded "
        "outcomes. Participants rated statements on a five-point scale, where one indicated strong "
        "disagreement and five indicated strong agreement. Table 5.1 summarises the items used "
        "during the session."
    )

    table = doc.add_table(rows=8, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "User testing question"
    for col, label in zip(range(2, 7), ["1", "2", "3", "4", "5"]):
        hdr[col].text = label

    questions = [
        "The system allows tourists to locate species and park information before a visit.",
        "The map helps users understand gates, sectors, and nearby points of interest.",
        "Tourists can submit or review wildlife sightings in an understandable way.",
        "Guides can access schedules and operational tools relevant to their treks.",
        "IT managers can view analytics summaries and operational status indicators.",
        "Tour-help responses stay relevant to Bwindi and park safety themes.",
        "The application remains usable when the network connection is interrupted briefly.",
    ]
    for idx, q in enumerate(questions, start=1):
        row = table.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = q

    cap = doc.add_paragraph("Table 5.1: User testing questionnaire for SIGTS (Likert scale 1–5).")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, 11, italic=True)

    body(
        "Unit testing was conducted on discrete server routes and client helpers. Individual "
        "components—such as authentication middleware, sightings validation, analytics report "
        "generation, and offline synchronisation queues—were exercised in isolation before being "
        "combined in end-to-end scenarios. Defects discovered at this stage were corrected prior "
        "to integration testing across the full stack."
    )

    heading("5.4 System Validation", 2)
    body(
        "Validation focused on whether data accepted by the system conformed to expected formats "
        "and business rules. Input checks were applied to registration fields, sighting coordinates, "
        "feedback ratings, and administrative forms. Server-side validation complemented client "
        "prompts so that incomplete or out-of-range values were rejected with clear messages. "
        "Retrieved records were compared against seeded catalogue content to ensure that tourists "
        "and guides received consistent information during demonstrations."
    )

    doc.add_page_break()

    heading("CHAPTER SIX: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS", 1)
    body(
        "6.0 Introduction\n"
        "This chapter closes the technical account of the Smart Information Guide Tour System. "
        "It restates what the project achieved, notes operational challenges observed during "
        "development and testing, acknowledges limitations of the study, and offers "
        "recommendations for future deployment at Bwindi Impenetrable National Park."
    )
    heading("6.1 Achievements", 2)
    body(
        "The project delivered a working prototype that supports tourist discovery of biodiversity "
        "and cultural content, map-based orientation within the park context, guide-facing "
        "workflow tools, and administrative analytics grounded in live database records. Offline "
        "caching and queued synchronisation were implemented so that brief loss of connectivity "
        "in the field does not prevent users from reading cached material or deferring sightings "
        "until the device reconnects. Security controls—including role separation, token-based "
        "sessions, and production configuration checks—were incorporated to align the build with "
        "responsible pilot deployment."
    )
    heading("6.2 Challenges", 2)
    body(
        "Several challenges emerged during implementation. Reliable connectivity inside and around "
        "the park remains uneven; features that depend on continuous access to the application "
        "programming interface degrade gracefully but cannot replace on-site ranger briefings. "
        "Maintaining accurate species and narrative content requires ongoing collaboration with "
        "Uganda Wildlife Authority and community partners. Optional tour-help capabilities depend "
        "on external language-model services when enabled; without them, the system relies on "
        "rule-based responses that are dependable yet less conversational. Server availability "
        "remains essential for authentication, synchronisation, and analytics used by IT staff."
    )
    heading("6.3 Limitations of the study", 2)
    body(
        "The study was constrained by time, equipment, and access to field respondents. User "
        "testing drew on a modest sample drawn largely from academic and demonstration settings "
        "rather than a full season of tourist traffic across all four sectors of the park. "
        "Hardware for load testing at scale was limited, so concurrent-usage behaviour under "
        "peak trekking periods was inferred rather than measured exhaustively. Some specification "
        "items—such as native mobile packaging, SMS-based multi-factor authentication, and "
        "full ranger-dispatch integration—were documented as future work rather than completed "
        "within the project window."
    )
    heading("6.4 Recommendations", 2)
    body(
        "Park management and the development team are advised to pilot SIGTS with trained guides "
        "and a controlled group of visitors before wider release. Short training sessions should "
        "cover map use, sighting etiquette, cultural content sensitivity, and realistic "
        "expectations for tour-help responses. Content stewards should establish a schedule for "
        "reviewing species records, FAQs, and safety tips. Production hosting should use strong "
        "secrets management, routine database backups, and monitoring of server health. "
        "Continued investment in connectivity at sector offices would increase the value of "
        "real-time analytics for operations staff."
    )
    heading("6.5 Conclusion", 2)
    body(
        "The Smart Information Guide Tour System was undertaken to improve how visitors prepare "
        "for and experience Bwindi Impenetrable National Park while giving guides and administrators "
        "shared tools for information delivery and monitoring. The implemented prototype "
        "demonstrates that a web-centred architecture with spatial data, offline tolerance, and "
        "role-aware interfaces can support this objective in a demonstrable form. The technologies "
        "chosen proved adequate for academic delivery and pilot demonstration, provided that "
        "content governance and infrastructure maintenance remain active responsibilities of the "
        "park. Further refinement—particularly in mobile distribution, extended field evaluation, "
        "and integration with official permit and ranger systems—will determine whether SIGTS "
        "transitions from a project outcome to a sustained operational service."
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
