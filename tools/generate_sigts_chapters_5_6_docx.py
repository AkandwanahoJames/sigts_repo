"""Generate SIGTS report chapters 5 and 6 as a standalone Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:\Projects\SIGTS\DOCS\SIGTS_Chapters_5_and_6_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(12)
    return p


def add_table_user_testing(doc):
    table = doc.add_table(rows=8, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    labels = ["No.", "Evaluation criterion", "1", "2", "3", "4", "5"]
    for i, lab in enumerate(labels):
        hdr_cells[i].text = lab
    rows_data = [
        (
            "1",
            "Tourists can register, sign in, and reach role-appropriate home screens",
        ),
        (
            "2",
            "Wildlife catalogue, map, and cultural content load inside the park boundary",
        ),
        (
            "3",
            "Guides can record sightings and view assigned tour information",
        ),
        (
            "4",
            "IT staff can access predictive analytics and operational summaries",
        ),
        (
            "5",
            "Tour help responds with park-grounded guidance when connectivity allows",
        ),
        (
            "6",
            "Offline queue retains sightings and syncs after reconnect",
        ),
        (
            "7",
            "Park access rules block off-site use while preserving cached content on site",
        ),
    ]
    for r_idx, (num, crit) in enumerate(rows_data, start=1):
        row = table.rows[r_idx].cells
        row[0].text = num
        row[1].text = crit
        for c in range(2, 7):
            row[c].text = ""
    cap = doc.add_paragraph("Table 5.1: User evaluation criteria for SIGTS (rating scale 1–5).")
    cap.paragraph_format.space_before = Pt(6)
    for run in cap.runs:
        run.font.size = Pt(11)
        run.italic = True


def build():
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(1)
    sect.bottom_margin = Inches(1)
    sect.left_margin = Inches(1.25)
    sect.right_margin = Inches(1.25)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Smart Information Guide Tour System (SIGTS)\n"
        "Chapters 5 and 6 — Implementation, Testing, and Project Closure"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    add_body(
        doc,
        "The following sections adapt the report structure used for the Group 34 end-of-year "
        "project template to the SIGTS deliverable developed for Bwindi Impenetrable National Park. "
        "Wording describes the implemented SIGTS prototype and its evaluation; it does not reproduce "
        "text from other student projects.",
    )

    # —— CHAPTER FIVE ——
    add_heading(doc, "CHAPTER FIVE: SYSTEM IMPLEMENTATION, TESTING AND VALIDATION", level=1)

    add_heading(doc, "5.0 Introduction", level=2)
    add_body(
        doc,
        "This chapter records how the approved SIGTS design was translated into a working "
        "application, how core modules behaved under test, and how entered data were checked against "
        "expected formats. The description covers the tourist, guide, and IT manager roles served by "
        "the same codebase.",
    )

    add_heading(doc, "5.1 System Implementation", level=2)
    add_body(
        doc,
        "Implementation followed the modular layout defined during analysis: a browser-based client, "
        "a Node.js API layer, and a PostgreSQL database with PostGIS extensions for park boundaries "
        "and location queries. Work was organised so that authentication, content services, field "
        "operations, and analytics could be delivered incrementally and demonstrated on local servers "
        "before deployment.",
    )

    add_heading(doc, "5.1.1 System technologies", level=3)
    add_body(
        doc,
        "The presentation tier uses HTML5, CSS, and client-side JavaScript organised into view, "
        "data-access, and manager modules. Leaflet provides interactive mapping. A service worker "
        "caches static assets and selected catalogue responses so that tourists and guides can continue "
        "browsing when connectivity drops inside the forest.",
    )
    add_body(
        doc,
        "The application tier runs on Node.js with Express. JSON Web Tokens secure authenticated "
        "requests; refresh rotation and idle-session checks reduce exposure from abandoned devices. "
        "Socket.IO supports optional real-time channels for tour coordination and IT monitoring. "
        "Public park content, sightings, tours, cultural narratives, geofence checks, and analytics "
        "are exposed through versioned REST endpoints under /api.",
    )
    add_body(
        doc,
        "PostgreSQL stores relational data for users, animals, locations, sightings, FAQs, safety "
        "guidance, and operational logs. PostGIS geometry types represent the park boundary and gate "
        "coordinates used during access validation. Migrations and seed scripts in the database folder "
        "keep schema changes reproducible across machines used by the project team.",
    )
    add_body(
        doc,
        "These tools were selected because they are openly documented, run on modest hardware, and "
        "match the skills available within the group. They also allow the same stack to be hosted on "
        "a single park server or split behind nginx when the park IT unit moves to production.",
    )

    add_heading(doc, "5.2 System Functionality", level=2)

    add_heading(doc, "5.2.1 Authentication and access control", level=3)
    add_body(
        doc,
        "All roles authenticate before protected screens load. Tourists and guides register with "
        "profile details; email verification and password reset flows are supported on the API. "
        "Role-based routing limits guides to shift and tour tools and confines IT managers to "
        "administration, analytics, and intranet features. Geofence middleware consults the visitor "
        "position—or an approved simulation flag during demonstrations—to permit full use only when "
        "the device reports coordinates inside the Bwindi boundary.",
    )

    add_heading(doc, "5.2.2 Wildlife catalogue, map, and tour guidance", level=3)
    add_body(
        doc,
        "Tourists browse species records enriched with conservation status, habitat notes, and "
        "stay-safe reminders drawn from seeded biodiversity content. The map view plots gates, "
        "viewpoints, and ranger posts; optional GPS updates refine nearby context. Ranked dashboard "
        "cards and a conversational tour-help panel answer common questions about treks, seasons, "
        "and park sectors when the language model endpoint is configured, otherwise a rule-based "
        "interpreter grounded on the same database tables supplies replies.",
    )

    add_heading(doc, "5.2.3 Cultural narratives and park information", level=3)
    add_body(
        doc,
        "Published cultural stories, FAQs, safety tips, and destination guides are retrieved from "
        "the server and cached on the device. IT staff verify and publish narratives through "
        "administrative routes before tourists see them. The information module therefore separates "
        "editorial control from public consumption in the same way the reference template separated "
        "administrator posts from patient-facing articles.",
    )

    add_heading(doc, "5.2.4 Sightings, tours, and analytics operations", level=3)
    add_body(
        doc,
        "Guides and tourists can submit wildlife sightings with species, location, and optional "
        "notes; rare events can raise alerts for staff review. Guide dashboards list scheduled tours, "
        "guest counts, and shift clock-in status. The IT predictive analytics workspace aggregates "
        "visitor flow, satisfaction scores, sightings trends, anomaly flags, exportable reports, and "
        "an operational summary endpoint that rolls up recent activity for park management meetings.",
    )

    add_heading(doc, "5.3 System Testing", level=2)
    add_body(
        doc,
        "Testing combined manual walkthroughs with module-level checks. Each sprint ended with "
        "scenarios executed on both localhost and a park-network test host so that role switching, "
        "hash routing, and offline queues were observed under realistic latency.",
    )

    add_heading(doc, "5.3.1 User testing", level=3)
    add_body(
        doc,
        "Representatives from the target user groups—prospective tourists, field guides, and IT "
        "support staff—were asked to rate statements on a five-point scale after guided tasks. "
        "Table 5.1 lists the criteria; numeric scores were recorded on paper during the session and "
        "averaged for the project log.",
    )
    add_table_user_testing(doc)

    add_heading(doc, "Unit testing", level=3)
    add_body(
        doc,
        "Individual API handlers were exercised with valid and invalid payloads before integration "
        "testing. Examples included sighting creation without mandatory animal identifiers, login "
        "with expired tokens, and analytics date ranges spanning empty tables. Front-end forms were "
        "checked for required fields and maximum lengths matching server validators.",
    )

    add_heading(doc, "5.4 System Validation", level=2)
    add_body(
        doc,
        "Validation focused on input integrity and business rules. Email and password fields enforce "
        "format and length limits; UUID keys tie sightings to catalogue species and map locations. "
        "Published cultural content cannot display until verification timestamps are set. Analytics "
        "exports reject unknown metric names. Together these checks ensure that data stored in "
        "PostgreSQL remain consistent with the field study requirements documented in Chapter Three.",
    )

  # —— CHAPTER SIX ——
    add_heading(doc, "CHAPTER SIX: SUMMARY, LIMITATIONS, CONCLUSION AND RECOMMENDATIONS", level=1)

    add_heading(doc, "6.0 Introduction", level=2)
    add_body(
        doc,
        "This chapter closes the SIGTS project by summarising outcomes, noting constraints encountered "
        "during development and field work, and stating recommendations for park adoption.",
    )

    add_heading(doc, "6.1 Achievements", level=2)
    add_body(
        doc,
        "The team delivered a single deployable system that unifies visitor education, guide "
        "operations, and IT oversight for Bwindi Impenetrable National Park. Tourists receive "
        "species information, map guidance, and safety copy without relying on paper leaflets alone. "
        "Guides gain a digital channel for sightings and tour preparation. IT managers access "
        "dashboards, predictive analytics, and backup-oriented operations tooling within the same "
        "security model. Offline caching and a retrying sync queue allow partial use when forest "
        "coverage interrupts mobile data.",
    )

    add_heading(doc, "6.2 Challenges", level=2)
    add_body(
        doc,
        "Field connectivity remains uneven across sectors; users near Buhoma may experience stable "
        "links while Rushaga trails lose signal for long periods. The prototype therefore depends on "
        "advance caching and cannot guarantee instant server confirmation for every sighting. "
        "Hosting also requires the API process and database service to remain available; park IT "
        "must plan power and backup for the server room. Configuring the external language model for "
        "tour help demands an API key and budget review before public rollout.",
    )

    add_heading(doc, "6.3 Limitations of the study", level=2)
    add_body(
        doc,
        "The study period allowed only a small number of guided user-test sessions, so statistical "
        "generalisation to all annual visitors is not claimed. Hardware for continuous GPS sampling "
        "on every tourist device was not issued; geofence accuracy therefore reflects browser "
        "location APIs rather than dedicated trackers. Some specification items—native mobile "
        "binaries, SMS-based MFA, and full offline map tile coverage—remain partial in the web "
        "prototype. Budget and time restricted extended trials with Uganda Wildlife Authority "
        "rangers beyond the project supervisors.",
    )

    add_heading(doc, "6.4 Recommendations", level=2)
    add_body(
        doc,
        "Park management should pilot SIGTS with one sector gate before park-wide promotion. "
        "Short training sessions for guides and front-desk staff will reduce support calls during "
        "the first trekking season. IT staff should run database migrations on a staging server, "
        "set production-grade secrets, and monitor the operational summary dashboard weekly. "
        "Future work may add SMS alerts for rare sightings, deeper elevation models on trails, and "
        "integration with Uganda Wildlife Authority permit systems when APIs become available.",
    )

    add_heading(doc, "6.5 Conclusions", level=2)
    add_body(
        doc,
        "SIGTS was built to narrow the information gap between visitors and park authorities at "
        "Bwindi Impenetrable National Park. The chosen web and database technologies proved adequate "
        "for a demonstrable system that respects role separation, location rules, and offline "
        "tolerance. Continued relevance depends on keeping species and safety content current, "
        "maintaining server infrastructure, and extending tests with a larger visitor sample. "
        "With those measures, the application can support safer treks, richer interpretation of "
        "biodiversity, and evidence-based decisions by park IT staff.",
    )

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
